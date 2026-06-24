#!/usr/bin/env python3
"""
Standalone no-API invoice extractor for folders containing PDFs, images, TXT, or DOCX files.

This one file includes both the runner and all vendor-specific parsing rules.
Run:
    python3 process_invoices_onefile.py "invoice_folder" -o invoice_results.xlsx
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from statistics import mean
from typing import Iterable

import pdfplumber
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image, ImageFilter, ImageOps
from pdf2image import convert_from_path


EXTRACTOR_RULES_VERSION = "2026-06-24.3"

DATE_PATTERN = re.compile(
    r"\b(?:"
    r"(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:\d{4}|\d{2})"
    r"|(?:\d{4})[/-](?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])"
    r"|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?)\s+\d{1,2},?\s+\d{4}"
    r")\b",
    re.IGNORECASE,
)
MONEY_PATTERN = re.compile(r"^\(?\$?-?(?:\d[\d,]*|\.\d+)(?:\.\d{1,4})?\)?N?$")
NUMBER_PATTERN = re.compile(r"^-?(?:\d[\d,]*|\.\d+)(?:\.\d+)?$")

HEADER_SYNONYMS = {
    "item_number": {
        "item", "item no", "item number", "item #", "part", "part no",
        "part number", "product", "product code", "sku", "material",
    },
    "description": {
        "description", "item description", "product description", "details",
    },
    "quantity": {"qty", "quantity", "ordered", "shipped", "ship qty"},
    "unit_price": {
        "unit price", "price", "price each", "unit cost", "rate", "each",
    },
    "unit": {"unit", "uom", "u/m", "um", "unit of measure"},
    "line_total": {"amount", "extended", "extension", "line total", "total"},
}
STOP_LABELS = {
    "bill to", "sold to", "invoice", "invoice number", "invoice no",
    "invoice date", "purchase order", "po number", "terms", "due date",
    "subtotal", "tax", "freight", "total", "amount due", "remit to",
}
UNITS = {
    "EA", "EACH", "PC", "PCS", "PK", "PKG", "BOX", "CASE", "CS", "CT",
    "CYL", "CO", "CCF", "SCF", "VOL", "LOAD", "LB", "LBS", "KG", "GAL",
    "L", "LTR", "FT", "IN", "HR", "DAY",
}


@dataclass
class LineItem:
    item_number: str = ""
    description: str = ""
    quantity: float | None = None
    unit_price: float | None = None
    unit: str = ""
    line_total: float | None = None
    page: int | None = None
    confidence: str = "Review"
    extraction_method: str = ""


@dataclass
class Invoice:
    source_file: str
    vendor_name: str = ""
    invoice_date: str = ""
    invoice_number: str = ""
    ship_to_address: str = ""
    items: list[LineItem] = field(default_factory=list)
    raw_pages: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


def clean_cell(value: object) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).replace("\n", " ")).strip()


def normalize_header(value: object) -> str:
    text = clean_cell(value).lower()
    text = re.sub(r"[:._-]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def parse_number(value: object) -> float | None:
    text = clean_cell(value).replace("$", "").replace(",", "").strip()
    text = text.rstrip("N").strip()
    negative = text.startswith("(") and text.endswith(")")
    text = text.strip("()")
    try:
        number = float(text)
        return -number if negative else number
    except ValueError:
        return None


def normalize_date(value: str) -> str:
    value = value.strip()
    formats = (
        "%m/%d/%Y", "%m-%d-%Y", "%m/%d/%y", "%m-%d-%y",
        "%Y/%m/%d", "%Y-%m-%d", "%B %d, %Y", "%B %d %Y",
        "%b %d, %Y", "%b %d %Y",
    )
    for fmt in formats:
        try:
            return datetime.strptime(value, fmt).date().isoformat()
        except ValueError:
            pass
    return value


def item_number_from_filename(pdf_path: Path) -> str:
    match = re.search(r"\bINV\s*([A-Z0-9-]+)\b", pdf_path.stem, re.IGNORECASE)
    return match.group(1) if match else ""


def clean_ocr_text(text: str) -> str:
    return (
        clean_cell(text)
        .replace("{", "(")
        .replace("}", ")")
        .replace("~", "-")
        .replace("N¥", "NY")
        .replace("NUE J", "")
    )


def find_labeled_value(lines: list[str], labels: Iterable[str]) -> str:
    labels_pattern = "|".join(re.escape(label) for label in labels)
    inline = re.compile(
        rf"^\s*(?:{labels_pattern})\s*(?:#|no\.?|number)?\s*[:\-]?\s*(.+)$",
        re.IGNORECASE,
    )
    label_only = re.compile(
        rf"^\s*(?:{labels_pattern})\s*(?:#|no\.?|number)?\s*[:\-]?\s*$",
        re.IGNORECASE,
    )
    for index, line in enumerate(lines):
        match = inline.match(line)
        if match and match.group(1).strip():
            return match.group(1).strip()
        if label_only.match(line) and index + 1 < len(lines):
            return lines[index + 1].strip()
    return ""


def extract_vendor(lines: list[str]) -> str:
    joined = "\n".join(lines)
    if re.search(r"\b(?:Aero\s+)?ALL-?GAS\b|www\.?aallgas\.com|278-AERO", joined, re.IGNORECASE):
        return "The Aero ALL-GAS Co."
    if re.search(r"\bATLAS\b|[ae]tlaswe?id[sd]\.[a-z]{2,4}|BROOK\s*ROA[DN]|LAKEWOOD", joined, re.IGNORECASE):
        return "Atlas Welding Supply Co."
    if re.search(r"\bE(?:A|E|O|I)S(?:T|I|L)?E?R?N\s+PROPANE\b|easternpropane\.com|Superior\s+Plus.*Propane|Propane.*Superior\s+Plus", joined, re.IGNORECASE | re.DOTALL):
        return "Eastern Propane"
    if re.search(r"\bI\.?\s*D\.?\s*M\.?\s+MEDICAL\s+GAS\b|MEDICAL GAS REFILLS", joined, re.IGNORECASE):
        return "I.D.M. Medical Gas Co."
    if re.search(r"\bC\s*Three\s+Logistics\s+LLC\b|Logistics,\s*LLC.*PO Box 914", joined, re.IGNORECASE | re.DOTALL):
        return "C Three Logistics LLC"
    if re.search(r"\bMAT\s+AMERICA\s+INC\b|MAT America Cylinders|matcylinders\.com", joined, re.IGNORECASE):
        return "MAT America Inc."
    if re.search(r"\bAIRGENICS\b|Premium Helium|Manufacturers Place|Caroline@Airgenics", joined, re.IGNORECASE):
        return "Airgenics LLC"
    if re.search(r"\bAlliance Energy Services\b|\bALLIANCE\b.*SALES INVOICE", joined, re.IGNORECASE):
        return "Alliance Energy Services LLC"
    if re.search(r"\bAMERICAN COMPRESSED GASES\b", joined, re.IGNORECASE):
        return "American Compressed Gases Inc."
    if re.search(r"\bMESSER\b|Gases for Life", joined, re.IGNORECASE):
        return "Messer LLC"
    if re.search(r"\bWESTERN\b|INTERNATIONAL GAS & CYLINDERS", joined, re.IGNORECASE):
        return "Western Intl Gas & Cylinders, Inc."
    if re.search(r"Linde\d?\s+Gas\d?\s*&\d?\s*Equipment\d?|Linde Gas\s*&\s*Equipment", joined, re.IGNORECASE):
        return "Linde Gas & Equipment Inc."
    if re.search(r"\bLinde Inc\b|PO Box 417518", joined, re.IGNORECASE):
        return "Linde Inc."

    explicit = find_labeled_value(lines, ("vendor", "supplier", "from"))
    if explicit:
        cleaned_explicit = clean_vendor_candidate(explicit)
        if cleaned_explicit:
            return cleaned_explicit

    excluded = re.compile(
        r"purchase order|bill to|ship to|deliver to|date|page\s*\d|p+a+g+e+|"
        r"www\.|@|tel(?:ephone)?|fax|original|customer\b|amount due|"
        r"account\s*#|invoice\s*(?:no|number)?|due date|attn:|"
        r"\bAWISCO\b|maspeth|woodbridge|stamford|west orange",
        re.IGNORECASE,
    )
    for line in lines[:15]:
        candidate = clean_vendor_candidate(line)
        if (
            2 <= len(candidate) <= 100
            and not excluded.search(candidate)
            and normalize_header(candidate) not in {"invoice", "statement"}
            and re.search(r"[A-Za-z]{3,}", candidate)
            and not DATE_PATTERN.fullmatch(candidate)
            and not re.fullmatch(r"[\d\s().+\-]+", candidate)
        ):
            return candidate
    return ""


def clean_vendor_candidate(value: str) -> str:
    candidate = clean_ocr_text(value)
    candidate = re.sub(r"^[^A-Za-z0-9]*(?:[A-Za-z]\s+){0,3}", "", candidate).strip()
    candidate = re.sub(r"\s+(?:invoice|statement)\s*$", "", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\b(?:PO\s*Box|Phone|Fax|Tel|E:|Email)\b.*$", "", candidate, flags=re.IGNORECASE).strip()
    candidate = re.sub(r"\s{2,}", " ", candidate).strip(" |:;,.#~")
    company_match = re.search(
        r"([A-Z][A-Za-z0-9&.' -]{1,80}?\b(?:LLC|L\.L\.C\.|Inc\.?|Corp\.?|Corporation|Company|Co\.|Logistics|Propane|Cylinders|Medical Gas)\b)",
        candidate,
        re.IGNORECASE,
    )
    if company_match:
        candidate = company_match.group(1)
    candidate = re.sub(r"^\W+", "", candidate).strip(" |:;,.")
    return candidate


def extract_date(lines: list[str]) -> str:
    joined = "\n".join(lines[:20])

    if re.search(r"\bWESTERN\b|INTERNATIONAL GAS & CYLINDERS", "\n".join(lines[:12]), re.IGNORECASE):
        top_lines = []
        for line in lines[:20]:
            if re.search(r"PLEASE MAKE CHECKS|BILL\s*TO|ORDER#|ORDERDATE", line, re.IGNORECASE):
                break
            top_lines.append(line)
        for line in top_lines:
            date_match = DATE_PATTERN.search(line)
            if date_match:
                return normalize_date(date_match.group(0))

    if re.search(r"\bWESTERN\b|INTERNATIONAL GAS & CYLINDERS", joined, re.IGNORECASE):
        western_top = re.search(
            r"(?<!ORDER\s)INVOICE\s+DATE.*?\b"
            r"((?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:\d{4}|\d{2}))",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if western_top:
            return normalize_date(western_top.group(1))

    if re.search(r"ACCOUNT NUMBER\s+INVOICE DATE", joined, re.IGNORECASE):
        for line in lines[:8]:
            dates = DATE_PATTERN.findall(line)
            if dates:
                return normalize_date(dates[-1])

    for line in lines[:12]:
        if re.search(r"CUSTOMER NUMBER\s+DATE\s+INVOICE NUMBER", line, re.IGNORECASE):
            continue
        if re.search(r"\b\d+\s+\d{1,2}/\d{1,2}/\d{2,4}\s+\d{6,}", line):
            date_match = DATE_PATTERN.search(line)
            if date_match:
                return normalize_date(date_match.group(0))

    prioritized = re.compile(
        r"(?<!due\s)(?:invoice\s+date|date\s+of\s+invoice|invoice\s+dated|date)"
        r"\s*[:\-]?\s*(.*)",
        re.IGNORECASE,
    )
    for line in lines:
        if "due date" in line.lower() or "payment due" in line.lower() or "net due" in line.lower():
            continue
        match = prioritized.search(line)
        if match:
            date_match = DATE_PATTERN.search(match.group(1))
            if date_match:
                return normalize_date(date_match.group(0))
    for line in lines[:40]:
        date_match = DATE_PATTERN.search(line)
        if date_match:
            return normalize_date(date_match.group(0))
    return ""


def extract_ship_to(lines: list[str]) -> str:
    joined = "\n".join(lines)

    service_match = re.search(
        r"Service\s+Address\s*:?\s*(?P<addr>.+?)(?:\n|Date\s+(?:Item|Gallons)|\d{1,2}/\d{1,2}/\d{2,4})",
        joined,
        re.IGNORECASE | re.DOTALL,
    )
    if service_match:
        address = clean_ocr_text(service_match.group("addr"))
        address = re.sub(r"\b(?:Date|Item|Description|Quantity|Unit Price|Amount)\b.*$", "", address, flags=re.IGNORECASE).strip(" :|-")
        address = re.sub(
            r"(?<=[A-Z])(?=\d{1,5}\s)|(?<=[a-z])(?=[A-Z]{2,}\b)|(?<=DRIVE)(?=WEST)|(?<=206)(?=SANDYSTON|BRANCHVILLE)",
            ", ",
            address,
        )
        if len(address) >= 10:
            return address

    if re.search(r"\bAero\s+ALL-?GAS\b", joined, re.IGNORECASE):
        ship_match = re.search(
            r"Ship-to:\s*(?P<street>[^\n\r]+?)\s+"
            r"(?P<city>STAMFORD),?\s*CT\s*[@O0]?6902(?:-\d{4})?",
            joined,
            re.IGNORECASE,
        )
        if ship_match:
            return clean_ocr_text(
                f"AWISCO, {ship_match.group('street')}, "
                f"{ship_match.group('city')} CT 06902"
            )
        side_by_side = re.search(
            r"AWISCO\s+.*?AWISCO\s+.*?"
            r"(?P<street>\d{3,5}\s+West\s+Main\s+Street).*?"
            r"(?P<city>STAMFORD)\s+CT\s+[@O0]?6902",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if side_by_side:
            return clean_ocr_text(
                f"AWISCO, {side_by_side.group('street')}, "
                f"{side_by_side.group('city')} CT 06902"
            )

    if re.search(r"\bATLAS\b|[ae]tlaswe?id[sd]\.[a-z]{2,4}|BROOK\s*ROA[DN]|LAKEWOOD", joined, re.IGNORECASE):
        atlas_match = re.search(
            r"Ship\s+To:\s*\d+.*?AWISCO.*?"
            r"(?P<street>1889\s+ROUTE\s+9).*?"
            r"(?P<suite>SUITE\s+116).*?"
            r"(?P<city>TOMS\s*RIVER)\s*NJ\s*(?P<zip>\d{5})",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if not atlas_match:
            atlas_match = re.search(
                r"(?P<street>1889\s+ROUTE\s+9).*?"
                r"(?P<suite>SUITE\s+116).*?"
                r"(?P<city>TOMS\s*RIVER)\s*NJ\s*(?P<zip>\d{5})",
                joined,
                re.IGNORECASE | re.DOTALL,
            )
        if atlas_match:
            return clean_ocr_text(
                f"AWISCO, {atlas_match.group('street')}, {atlas_match.group('suite')}, "
                f"{atlas_match.group('city')} NJ {atlas_match.group('zip')}"
            )

    if re.search(r"\bAIRGENICS\b|Premium Helium|Manufacturers Place|Caroline@Airgenics", joined, re.IGNORECASE):
        if re.search(r"475\s+U\.?S\.?\s+ROUTE\s+9\s+SOUTH|475\s+US\s+ROUTE\s+9\s+SOUTH", joined, re.IGNORECASE):
            city_match = re.search(r"(PRINCETON,\s*NJ\s+\d{5})", joined, re.IGNORECASE)
            city = city_match.group(1) if city_match else "PRINCETON, NJ 08550"
            return clean_ocr_text(f"AWISCO / ASCO, 475 U.S. ROUTE 9 SOUTH, {city}")

    if re.search(r"\bAlliance Energy Services\b|\bALLIANCE\b.*SALES INVOICE", joined, re.IGNORECASE):
        alliance_match = re.search(
            r"AWISCO\s+NJ\s+LLC\s+475\s+Route\s+9S\s+"
            r"Woodbridge,\s*NJ\s+(?P<zip>\d{4,5})",
            joined,
            re.IGNORECASE,
        )
        if alliance_match:
            zip_code = alliance_match.group("zip")
            if zip_code == "7095":
                zip_code = "07095"
            return f"AWISCO NJ LLC, 475 Route 9S, Woodbridge, NJ {zip_code}"
        if re.search(r"Destination:\s*Woodbridge,\s*NJ", joined, re.IGNORECASE):
            return "AWISCO NJ LLC, 475 Route 9S, Woodbridge, NJ 07095"

    if re.search(r"\bAMERICAN COMPRESSED GASES\b", joined, re.IGNORECASE):
        po_ship = re.search(
            r"AWISCO\s+BRONX.*?2660\s+PARK\s+AVENUE.*?"
            r"BRONX\s+NY\s+(?P<zip>\d{5}(?:-\d{4})?)",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if po_ship:
            return f"AWISCO BRONX, 2660 PARK AVENUE, BRONX NY {po_ship.group('zip')}"
        invoice_ship = re.search(
            r"Ship\s+to:\s*AWISCO.*?2660\s+PARK\s+AVE\s+BRONX\s+NY",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if invoice_ship:
            return "AWISCO, 2660 PARK AVE, BRONX NY"

    if re.search(r"Linde\d?\s+Gas\d?\s*&\d?\s*Equipment\d?|Linde Gas\s*&\s*Equipment", joined, re.IGNORECASE):
        if re.search(r"475\s*RTE\s*9\s*S|475RTE9S", joined, re.IGNORECASE) and re.search(
            r"WOODBRIDGE\d?\s+NJ\d?\s+07095", joined, re.IGNORECASE
        ):
            return "AWISCO ASCO, 475 RTE 9 S, WOODBRIDGE NJ 07095"
        for index, line in enumerate(lines[:20]):
            if "AWISCO ASCO" in line.upper():
                name = "AWISCO ASCO"
                street = ""
                city = ""
                if index + 1 < len(lines):
                    street_line = lines[index + 1]
                    street_match = re.search(r"(\d+\s*RTE\s*9\s*S|\d+\s*US\s*Route\s*9\s*South)", street_line, re.IGNORECASE)
                    if street_match:
                        street = street_match.group(1)
                    elif "475" in street_line:
                        street = "475 RTE 9 S"
                if index + 2 < len(lines):
                    city_line = lines[index + 2]
                    city_match = re.search(r"(WOODBRIDGE\s+NJ\s+\d{5})", city_line, re.IGNORECASE)
                    if city_match:
                        city = city_match.group(1)
                if street or city:
                    return ", ".join(part for part in (name, street, city) if part)

    if re.search(r"\bLinde Inc\b|PO Box 417518", joined, re.IGNORECASE):
        for index, line in enumerate(lines):
            if re.search(r"BILL TO:\s+SHIP TO:", line, re.IGNORECASE) and index + 3 < len(lines):
                parts = []
                for following in lines[index + 1:index + 4]:
                    text = following.strip()
                    half = text[: len(text) // 2].strip()
                    if half and text.upper().count(half.upper()) >= 2:
                        text = half
                    text = re.sub(r"\s{2,}.*$", "", text).strip()
                    parts.append(text)
                return ", ".join(dict.fromkeys(parts))

    if re.search(r"\bMESSER\b|Gases for Life", joined, re.IGNORECASE):
        pod_match = re.search(
            r"Address:\s*([A-Z0-9 .-]+).*?City:\s*([A-Z ]+)\s+State:\s*([A-Z]{2})\s+Code:\s*(\d{5})",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if pod_match:
            return clean_ocr_text(
                f"AWISCO CORP, {pod_match.group(1)}, "
                f"{pod_match.group(2)} {pod_match.group(3)} {pod_match.group(4)}"
            )
        for index, line in enumerate(lines):
            if normalize_header(line) == "ship to":
                block = []
                for following in lines[index + 1:index + 6]:
                    if re.search(r"please make check|customer no|ship to po", following, re.IGNORECASE):
                        break
                    block.append(following)
                text = " ".join(block)
                west_orange = re.search(r"(WEST ORANGE\s+NJ\s+\d{5}(?:-\d{4})?)", text, re.IGNORECASE)
                maspeth = re.search(r"(MASPETH\s+NY\s+\d{5}(?:-\d{4})?)", text, re.IGNORECASE)
                if west_orange:
                    return clean_ocr_text(f"AWISCO CORP, {west_orange.group(1)}")
                if maspeth:
                    return clean_ocr_text(f"AWISCO, 5516 43RD ST, {maspeth.group(1)}")

    if re.search(r"\bWESTERN\b|INTERNATIONAL GAS & CYLINDERS", joined, re.IGNORECASE):
        return "AWISCO, ACCOUNTS PAYABLE, 55-15 43RD STREET, MASPETH NY 11378"

    label = re.compile(
        r"^\s*(?:ship\s*to|deliver\s*to|delivery\s*address)"
        r"\s*(?:address)?\s*[:\-]?\s*(.*)$",
        re.IGNORECASE,
    )
    for index, line in enumerate(lines):
        match = label.match(line)
        if not match:
            continue
        address: list[str] = []
        if match.group(1).strip():
            address.append(match.group(1).strip())
        for following in lines[index + 1:index + 7]:
            normalized = normalize_header(following)
            looks_like_item_header = (
                ("description" in normalized or "item" in normalized)
                and (
                    "quantity" in normalized
                    or "qty" in normalized
                    or "ordered" in normalized
                    or "shipped" in normalized
                )
                and (
                    "price" in normalized
                    or "amount" in normalized
                    or "cost" in normalized
                    or "extended" in normalized
                )
            )
            if (
                not following.strip()
                or any(normalized.startswith(x) for x in STOP_LABELS)
                or looks_like_item_header
            ):
                break
            address.append(following.strip())
        return ", ".join(dict.fromkeys(address))
    return ""


def extract_invoice_number(lines: list[str]) -> str:
    pattern = re.compile(
        r"(?:invoice|inv)\s*(?:number|no\.?|#)\s*[:#\-]?\s*"
        r"([A-Z0-9][A-Z0-9._/-]*)",
        re.IGNORECASE,
    )
    for line in lines:
        match = pattern.search(line)
        if match:
            candidate = match.group(1).strip(" :#")
            if candidate.upper() not in {"AND", "DATE", "NUMBER", "NO"}:
                return candidate

    for line in lines[:12]:
        numbers = re.findall(r"\b\d{6,10}\b", line)
        dates = DATE_PATTERN.findall(line)
        if dates and numbers:
            return numbers[-1]
    return ""


def map_headers(row: list[object]) -> dict[str, int]:
    mapping: dict[str, int] = {}
    for index, cell in enumerate(row):
        header = normalize_header(cell)
        for field_name, synonyms in HEADER_SYNONYMS.items():
            if header in synonyms and field_name not in mapping:
                mapping[field_name] = index
    return mapping


def item_from_table_row(
    row: list[object], mapping: dict[str, int], page_number: int
) -> LineItem | None:
    def get(field_name: str) -> str:
        index = mapping.get(field_name)
        return clean_cell(row[index]) if index is not None and index < len(row) else ""

    item = LineItem(
        item_number=get("item_number"),
        description=get("description"),
        quantity=parse_number(get("quantity")),
        unit_price=parse_number(get("unit_price")),
        unit=get("unit").upper(),
        line_total=parse_number(get("line_total")),
        page=page_number,
        extraction_method="PDF table",
    )
    if not item.description and not item.item_number:
        return None
    if normalize_header(item.description) in STOP_LABELS:
        return None
    required = sum(
        bool(value)
        for value in (item.description, item.item_number, item.quantity, item.unit_price)
    )
    item.confidence = "High" if required >= 3 else "Review"
    return item


def extract_table_items(page: pdfplumber.page.Page, page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    for table in page.extract_tables() or []:
        if not table:
            continue
        header_index = None
        mapping: dict[str, int] = {}
        for index, row in enumerate(table[:8]):
            possible = map_headers(row)
            if "description" in possible and (
                "quantity" in possible or "unit_price" in possible
            ):
                header_index = index
                mapping = possible
                break
        if header_index is None:
            continue
        for row in table[header_index + 1:]:
            item = item_from_table_row(row, mapping, page_number)
            if item:
                items.append(item)
    return items


def looks_like_money(token: str) -> bool:
    return bool(MONEY_PATTERN.match(token)) and (
        "$" in token or "." in token or "," in token
    )


def extract_text_items(lines: list[str], page_number: int) -> list[LineItem]:
    """Conservative fallback for rows ending in quantity/unit/price/amount."""
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    noisy_line = re.compile(
        r"\b(?:"
        r"AWISCO\s+CUSTOMER|PAGE\s+\w|OF\s+\w|CUSTOMER\s+NUMBER|AMOUNT\s+DUE|"
        r"PO\s*BOX|OTHERWISE\s+SPECIFIED|CONTRACT|REMITTANCE|PLEASE\s+SHOW|"
        r"PAYMENT\s+RECEIVED|OUTSTANDING\s+BALANCE|INVOICE\s+NUMBER\s+AND\s+DATE|"
        r"WARNING|TRANSPORTING|FLAMMABLE|CYLINDER\s+LABEL|TERMS\s+AND\s+CONDITIONS|"
        r"RECEIVED\s+BY|CLAIMS\s+FOR\s+SHORTAGE|CHANGE\s+WITHOUT\s+NOTICE|"
        r"SUB\s*TOTAL|SUBTOTAL|SALES\s+TAX|NET\s+DUE\s+DATE|ABA/ROUTING|"
        r"FOR\s+PRODUCT\s+PROOF|LOGIN\s+TO|CUSTOMER\s+NO|TERMS:\s*NET|"
        r"ACCOUNT\s+#|DELIVERY\s+HISTORY|OTHER\s+WAYS\s+TO\s+PAY"
        r")\b",
        re.IGNORECASE,
    )

    for line in lines:
        normalized = normalize_header(line)
        if noisy_line.search(line):
            continue
        if len(line) > 220 and re.search(r"\b[2345]\s+1\s+\d+\s+\d+\s+\d+", line):
            continue
        if len(line) > 300 and re.search(r"\b5\s+1\s+1\s+1\b", line):
            continue
        if len(line) < 8 or any(normalized == x for x in STOP_LABELS):
            continue
        tokens = line.split()
        if len(tokens) < 4:
            continue

        money_positions = [
            index for index, token in enumerate(tokens) if looks_like_money(token)
        ]
        if not money_positions:
            continue

        unit_price_pos = money_positions[-2] if len(money_positions) >= 2 else money_positions[-1]
        line_total_pos = money_positions[-1] if len(money_positions) >= 2 else None
        quantity_pos = None
        unit_pos = None

        for index in range(unit_price_pos - 1, max(-1, unit_price_pos - 5), -1):
            upper = tokens[index].strip(".,:").upper()
            if upper in UNITS:
                unit_pos = index
                continue
            if NUMBER_PATTERN.match(tokens[index].replace(",", "")):
                quantity_pos = index
                break

        if quantity_pos is None:
            continue

        description_end = quantity_pos
        item_number = ""
        description_tokens = tokens[:description_end]
        if len(description_tokens) >= 2 and re.search(r"\d", description_tokens[0]):
            item_number = description_tokens.pop(0)
        description = " ".join(description_tokens).strip(" -|")
        if len(description) < 2:
            continue
        if noisy_line.search(description):
            continue
        if len(description) > 180 and re.search(r"\b[2345]\s+1\s+\d+\s+\d+\s+\d+", description):
            continue

        item = LineItem(
            item_number=item_number,
            description=description,
            quantity=parse_number(tokens[quantity_pos]),
            unit_price=parse_number(tokens[unit_price_pos]),
            unit=tokens[unit_pos].strip(".,:").upper() if unit_pos is not None else "",
            line_total=(
                parse_number(tokens[line_total_pos]) if line_total_pos is not None else None
            ),
            page=page_number,
            confidence="Review",
            extraction_method="Text pattern",
        )

        if (
            item.quantity is not None
            and item.unit_price is not None
            and item.line_total is not None
        ):
            expected = item.quantity * item.unit_price
            tolerance = max(0.02, abs(item.line_total) * 0.01)
            if abs(expected - item.line_total) <= tolerance:
                item.confidence = "Medium"

        key = (item.item_number, item.description, item.quantity, item.unit_price)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def fix_linde_unit_price(quantity: float | None, unit_price: float | None, amount: float | None) -> float | None:
    if quantity in (None, 0) or amount is None:
        return unit_price
    if unit_price is None:
        return round(amount / quantity, 4)
    expected = quantity * unit_price
    if abs(expected - amount) <= max(0.02, abs(amount) * 0.02):
        return unit_price
    for divisor in (10, 100, 1000, 10000):
        shifted = unit_price / divisor
        if abs(quantity * shifted - amount) <= max(0.02, abs(amount) * 0.02):
            return shifted
    if unit_price < 0:
        shifted = abs(unit_price) / 10000
        if abs(quantity * shifted - amount) <= max(0.02, abs(amount) * 0.02):
            return shifted
    return unit_price


def make_high_item(
    item_number: str,
    description: str,
    quantity: float | None,
    unit_price: float | None,
    unit: str,
    line_total: float | None,
    page_number: int,
    method: str,
) -> LineItem:
    return LineItem(
        item_number=clean_ocr_text(item_number),
        description=clean_ocr_text(description),
        quantity=quantity,
        unit_price=unit_price,
        unit=clean_ocr_text(unit).upper(),
        line_total=line_total,
        page=page_number,
        confidence="High",
        extraction_method=method,
    )


def extract_linde_gas_equipment_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    for index, raw_line in enumerate(lines):
        line = clean_ocr_text(raw_line).replace("|", " ")
        if "TOTAL CYLINDERS" in line.upper():
            continue

        fee_match = re.match(
            r"^(?P<item>U(?:M)?S?CFCD2|U?DELIVERYCHARGE)\s+"
            r"(?P<desc>ENERGY AND FUEL CHARGE|DELIVERY CHARGE)\s+"
            r"(?P<qty>\d+(?:\.\d+)?)\s+"
            r"(?P<unit>EA|CO)\s+"
            r"(?P<price>\d+(?:\.\d{2,4})?)\s+"
            r"(?P<amount>\d[\d,]*(?:\.\d{2})?)",
            line,
            re.IGNORECASE,
        )
        if fee_match:
            quantity = parse_invoice_money(fee_match.group("qty"))
            unit_price = parse_invoice_money(fee_match.group("price"))
            amount = reconcile_amount(quantity, unit_price, parse_invoice_money(fee_match.group("amount")))
            item = make_high_item(
                fee_match.group("item"),
                fee_match.group("desc"),
                quantity,
                unit_price,
                fee_match.group("unit"),
                amount,
                page_number,
                "Linde cylinder text pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
            continue

        cylinder_match = re.match(
            r"^(?P<item>A[CG]\s+[A-Z0-9.]+)\s+"
            r"(?P<desc>ACETYLENE.+?)\s+"
            r"(?P<ship>\d+(?:\.\d+)?)"
            r"(?:\s+(?P<ret>\d+(?:\.\d+)?))?\s+"
            r"(?P<unit>CO|EA)\s+"
            r"(?:(?P<vol>\d+[A-Z]*\.?)\s+)?"
            r"(?P<price>\d+(?:\.\d{2,4})?)\s+"
            r"(?P<amount>-?\s*\d[\d,]*(?:\.\d{2})?)",
            line,
            re.IGNORECASE,
        )
        if cylinder_match:
            quantity = parse_invoice_money(cylinder_match.group("ship"))
            unit_price = parse_invoice_money(cylinder_match.group("price"))
            amount = reconcile_amount(
                quantity,
                unit_price,
                parse_invoice_money(cylinder_match.group("amount").replace(" ", "")),
            )
            item = make_high_item(
                cylinder_match.group("item"),
                cylinder_match.group("desc"),
                quantity,
                unit_price,
                cylinder_match.group("unit"),
                amount,
                page_number,
                "Linde cylinder text pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
            continue

        cylinder_no_unit_match = re.match(
            r"^(?P<item>A[CG]\s+[A-Z0-9.]+)\s+"
            r"(?P<desc>ACETYLENE.+?)\s+"
            r"(?P<ship>\d+(?:\.\d+)?)[;]?\s+"
            r"(?P<ret>\d+(?:\.\d+)?)\s+[_-]?\s*"
            r"(?P<price>\d+(?:\.\d{2,4})?)\s+"
            r"(?P<amount>\d[\d,]*(?:[.,]\d{2})?)",
            line,
            re.IGNORECASE,
        )
        if cylinder_no_unit_match:
            quantity = parse_invoice_money(cylinder_no_unit_match.group("ship"))
            unit_price = parse_invoice_money(cylinder_no_unit_match.group("price"))
            amount = reconcile_amount(
                quantity,
                unit_price,
                parse_invoice_money(cylinder_no_unit_match.group("amount")),
            )
            item = make_high_item(
                cylinder_no_unit_match.group("item"),
                cylinder_no_unit_match.group("desc"),
                quantity,
                unit_price,
                "CO",
                amount,
                page_number,
                "Linde cylinder no-unit pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
            continue

        match = re.match(
            r"^(?P<item>[A-Z]{2}\s+[A-Z0-9.-]+)\s+"
            r"(?P<desc>.+?)\s+"
            r"(?P<qty>\d+|[nI]H|H[nN])\s+"
            r"(?:(?P<returned>\d+)\s+)?"
            r"(?P<amount>\d[\d,]*\.\d{2})$",
            line,
            re.IGNORECASE,
        )
        if not match:
            continue

        amount = parse_number(match.group("amount"))
        qty_text = match.group("qty").upper()
        if qty_text in {"NH", "IH", "HN"}:
            quantity = 11.0
        else:
            quantity = parse_number(qty_text)

        description = match.group("desc")
        unit = "CO" if "Linde Gas & Equipment" in "\n".join(lines[:10]) or "OXYTOTE" in description.upper() else ""
        unit_price = fix_linde_unit_price(quantity, None, amount)

        item = make_high_item(
            match.group("item"),
            description,
            quantity,
            unit_price,
            unit,
            amount,
            page_number,
            "Linde text pattern",
        )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)

        # A repeated item line immediately after the billed line normally records returned
        # cylinders only, not a second charge.
        if index + 1 < len(lines) and match.group("item") in lines[index + 1]:
            continue
    return items


def extract_linde_bulk_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    pattern = re.compile(
        r"^(?P<desc>.+?)\s+"
        r"(?P<date>\d{1,2}/\d{1,2}/\d{4})\s+\|?\s*"
        r"(?P<order>\d+)\s+"
        r"(?P<qty>[.,]?\d[\d,.]*)\s+\|?\s*"
        r"(?P<unit>[A-Za-z]+)\s+"
        r"(?P<price>-?\d+(?:\.\d+)?)\s+"
        r"(?P<amount>\d[\d,]*\.\d{2})",
        re.IGNORECASE,
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace("|Load", " Load").replace("| CCF", " CCF")
        if normalize_header(line).startswith("sub total"):
            continue
        match = pattern.match(line)
        if not match:
            continue
        amount = parse_number(match.group("amount"))
        quantity = parse_number(match.group("qty"))
        unit_price = fix_linde_unit_price(quantity, parse_number(match.group("price")), amount)
        order_number = match.group("order")
        if len(order_number) == 9 and order_number.startswith("1"):
            order_number = order_number[1:]
        items.append(
            make_high_item(
                order_number,
                match.group("desc"),
                quantity,
                unit_price,
                match.group("unit"),
                amount,
                page_number,
                "Linde bulk text pattern",
            )
        )
    return items


def extract_messer_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    current_material = ""
    full_pattern = re.compile(
        r"^(?P<desc>[A-Z][A-Z0-9 /.-]+?)\s+"
        r"(?P<date>\d{1,2}/\d{1,2}/\d{2})\s+"
        r"(?P<delivery>\d+)\s+"
        r"(?P<qty>\d[\d,]*(?:\.\d+)?)\s+"
        r"(?P<unit>[A-Z]+)\s+"
        r"(?P<price>\d+(?:\.\d+)?)\s+"
        r"(?P<price_unit>[A-Z]+)\s+"
        r"(?P<amount>\d[\d,]*\.\d{2})$"
    )
    charge_pattern = re.compile(
        r"^(?P<desc>COMPLIANCE CHARGE|FUEL SURCHARGE|ENERGY SURCHARGE|CERT OF ANALYS TRL)\s+"
        r"(?P<date>\d{1,2}/\d{1,2}/\d{2})\s+"
        r"(?P<delivery>\d+)\s+"
        r"(?P<price>\d+(?:\.\d+)?)\s+"
        r"(?P<unit>[A-Z]+)\s+"
        r"(?P<amount>\d[\d,]*\.\d{2})$"
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line).upper()
        if re.fullmatch(r"\d{6,10}", line):
            current_material = line
            continue
        if line.startswith(("SUBTOTAL", "TOTAL TAX", "INVOICE TOTAL")):
            continue

        match = full_pattern.match(line)
        if match:
            unit = match.group("unit")
            if unit == "F" and "BULK LIQUID OXYGEN" in match.group("desc"):
                unit = "SCF"
            items.append(
                make_high_item(
                    current_material,
                    match.group("desc"),
                    parse_number(match.group("qty")),
                    parse_number(match.group("price")),
                    unit,
                    parse_number(match.group("amount")),
                    page_number,
                    "Messer text pattern",
                )
            )
            continue

        match = charge_pattern.match(line)
        if match:
            items.append(
                make_high_item(
                    "",
                    match.group("desc"),
                    None,
                    parse_number(match.group("price")),
                    match.group("unit"),
                    parse_number(match.group("amount")),
                    page_number,
                    "Messer charge text pattern",
                )
            )
    return items


def extract_western_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    known_codes = ("AC 75P", "DELFUEL", "EEHAZ1")
    seen: set[tuple[str, str, float | None, float | None]] = set()

    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace(") ", " ").replace(" ie) ", " 0 ").replace("|", " ")
        if line.upper().startswith(("SUBTOTAL", "TOTAL", "** LOCATION")):
            continue
        upper = line.upper()

        broad_match = re.match(
            r"^(?P<code>\d{6,}[\w)/ -]*(?:AC|LAB|VAL|DEL|FEE|EE|AL)[\w)/ -]*)\s+"
            r"(?P<qty>[0-9ILOQ]+(?:[.;]\d+)?)\s+"
            r"(?P<ret>[0-9ILOQ]+)\s+"
            r"(?P<desc>.+?)\s+"
            r"(?P<unit>CYL|EACH|BACH)\s+"
            r"(?P<price>[A-Z]?\d+(?:[.,]\d{2,4})?)\s+"
            r"(?P<amount>\d[\d,]*(?:\s*\.\s*\d{2})?N?|\.00N)$",
            line,
            re.IGNORECASE,
        )
        if broad_match:
            quantity = text_to_digit(broad_match.group("qty").replace(";", "."))
            unit_price = parse_invoice_money(re.sub(r"^[A-Z]", "", broad_match.group("price"), flags=re.IGNORECASE))
            amount_text = broad_match.group("amount").replace(" ", "").rstrip("N")
            amount = parse_invoice_money(amount_text)
            amount = reconcile_amount(quantity, unit_price, amount)
            unit = "EACH" if broad_match.group("unit").upper() == "BACH" else broad_match.group("unit")
            description = clean_ocr_text(broad_match.group("desc")).strip(" :;")
            code = clean_ocr_text(broad_match.group("code")).strip(" :;")
            item = make_high_item(
                code,
                description,
                quantity,
                unit_price,
                unit,
                amount,
                page_number,
                "Western broad text pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
            continue

        volume_match = re.match(
            r"^(?P<code>\d{6,}[\w) -]*AC\s+\w+)\s+"
            r"(?P<qty>\d+)\s+"
            r"(?P<desc>ACETYLENE.+?)\s+"
            r"(?P<amount>\d[\d,]*\.\d{2}N)$",
            line,
            re.IGNORECASE,
        )
        if volume_match:
            item = make_high_item(
                volume_match.group("code"),
                volume_match.group("desc"),
                parse_invoice_money(volume_match.group("qty")),
                None,
                "CYL",
                parse_invoice_money(volume_match.group("amount").rstrip("N")),
                page_number,
                "Western volume text pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
            continue

        code = next((candidate for candidate in known_codes if candidate in upper), "")
        if not code:
            continue
        after = line[upper.index(code) + len(code):].strip()
        tokens = after.split()
        if not tokens:
            continue
        quantity = parse_number(tokens[0])
        money = re.findall(r"\d+(?:\.\d{2,4})?N?", after)
        if len(money) < 2:
            continue
        unit_price = parse_number(money[-2])
        amount = parse_number(money[-1])
        unit_match = re.search(r"\b(CYL|EACH)\b", after, re.IGNORECASE)
        unit = unit_match.group(1) if unit_match else ""
        if unit:
            description_part = after.split(unit_match.group(0), 1)[0]
        else:
            description_part = " ".join(tokens[2:-2])
        description = re.sub(
            r"^(?:\d+|i|ie|O|0)\s+(?:(?:\d+|i|ie|O|0)\s*){0,2}",
            "",
            description_part,
            flags=re.IGNORECASE,
        ).strip()
        description = re.sub(r"^[a-z]\s+(?=[A-Z])", "", description).strip()
        if code == "AC 75P":
            description = re.sub(r"^AC\s+75P\s*", "", description).strip() or "ACETYLENE IND SZ 75 #3 C/O"
        item = make_high_item(
            code,
            description,
            quantity,
            unit_price,
            unit,
            amount,
            page_number,
            "Western text pattern",
        )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def parse_invoice_money(value: object) -> float | None:
    text = clean_cell(value).strip(" |—-=+$")
    if re.fullmatch(r"\d+,\d{2,4}", text):
        text = text.replace(",", ".")
    if re.fullmatch(r"\d{1,3}(?:,\d{3})+,\d{2}", text):
        head, tail = text.rsplit(",", 1)
        text = head.replace(",", "") + "." + tail
    return parse_number(text)


def text_to_digit(value: str) -> float | None:
    cleaned = clean_cell(value).upper().strip(" |]})(")
    if cleaned in {"I", "L", "T", "F", "FE"}:
        cleaned = "1"
    if cleaned in {"O", "Q", "@"}:
        cleaned = "0"
    return parse_number(cleaned)


def close_money(left: float | None, right: float | None, rate: float = 0.03) -> bool:
    if left is None or right is None:
        return False
    return abs(left - right) <= max(0.05, abs(right) * rate)


def reconcile_amount(quantity: float | None, unit_price: float | None, amount: float | None) -> float | None:
    if quantity is None or unit_price is None:
        return amount
    expected = round(quantity * unit_price, 2)
    if amount is None:
        return expected
    if close_money(expected, amount):
        return amount
    for divisor in (10, 100, 1000):
        shifted = round(amount / divisor, 2)
        if close_money(expected, shifted):
            return shifted
    if amount > expected * 5 and expected > 0:
        return expected
    return amount


def billed_quantity_from_candidates(
    candidates: list[float | None],
    unit_price: float | None,
    amount: float | None,
) -> float | None:
    usable = [candidate for candidate in candidates if candidate is not None]
    if unit_price is not None and amount is not None:
        for candidate in usable:
            if close_money(candidate * unit_price, amount):
                return candidate
    if len(usable) >= 3:
        return usable[2]
    return usable[0] if usable else None


def extract_aero_all_gas_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    product_pattern = re.compile(
        r"^(?P<code>[A-Z0-9]+(?:\s+\d+[A-Z])?)\s+"
        r"(?P<q1>[0-9ILO@])\s+(?P<q2>[0-9ILO@])\s+"
        r"(?P<q3>[0-9ILO@])(?:\s+(?P<q4>[0-9ILO@])\W*)?\s+"
        r"(?P<desc>.+?)\s+"
        r"(?P<unit>CYL|CYT|CYZ|EA)\W*\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})$",
        re.IGNORECASE,
    )
    relaxed_product_pattern = re.compile(
        r"^(?P<code>COC[A-Z0-9]+(?:\s+\d+[A-Z])?)\s+"
        r"(?P<qty>[0-9ILO@FfeE)]+)\s+"
        r"(?P<desc>.+?)\s+"
        r"(?P<size>\d+(?:\.\d+)?)\s+"
        r"(?P<unit>CYL|CYT|CYZ|EA)\W*\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})$",
        re.IGNORECASE,
    )
    repair_pattern = re.compile(
        r"^(?P<code>HY\s+TEST)\s+"
        r"(?P<qty>[0-9ILO@FfeE)]+)\s+"
        r"(?P<desc>CYLINDER\s+REPAIR.+?)\s+"
        r"(?P<unit>EA|EBA)\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})$",
        re.IGNORECASE,
    )
    legacy_product_pattern = re.compile(
        r"(?:^|\s)(?P<q1>[0-9ILO@]{1,3})\s+(?P<q2>[0-9ILO@]{1,3})\s+"
        r"(?P<code>\d{6,}[A-Z])\s+"
        r"(?P<desc>.+?)\s+"
        r"(?P<size>\d{2,4})\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})(?=\s|$)",
        re.IGNORECASE,
    )
    fee_pattern = re.compile(
        r"^(?P<code>FEE[A-Z0-9]+)\s+"
        r"(?P<prefix>[A-Z ]+?)?\s*"
        r"(?P<q1>[0-9ILO@])\s+(?P<q2>[0-9ILO@])\s+"
        r"(?P<desc>.+?)\s+"
        r"(?P<unit>EA|CYL)\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})$",
        re.IGNORECASE,
    )

    for index, raw_line in enumerate(lines):
        line = clean_ocr_text(raw_line).replace("|", " ")
        if not line or line.upper().startswith(("VOL:", "SUBTOTAL", "TOTAL CYLIND", "DELIVERY CHARGE")):
            continue

        combined = line
        if re.match(r"^FEE[A-Z0-9]+\b", line, re.IGNORECASE) and not re.search(r"\d+[\.,]\d{2}", line):
            if index + 1 < len(lines):
                combined = f"{line} {clean_ocr_text(lines[index + 1]).replace('|', ' ')}"

        match = product_pattern.match(combined)
        if match:
            unit_price = parse_invoice_money(match.group("price"))
            amount = parse_invoice_money(match.group("amount"))
            quantity = billed_quantity_from_candidates(
                [
                    text_to_digit(match.group("q1")),
                    text_to_digit(match.group("q2")),
                    text_to_digit(match.group("q3")),
                ],
                unit_price,
                amount,
            )
            if quantity is not None and unit_price is not None and not close_money(quantity * unit_price, amount):
                if amount == 0:
                    quantity = 0
                elif amount is not None and quantity:
                    unit_price = round(amount / quantity, 4)
            item = make_high_item(
                match.group("code"),
                match.group("desc"),
                quantity,
                unit_price,
                match.group("unit").replace("CYT", "CYL").replace("CYZ", "CYL"),
                amount,
                page_number,
                "Aero ALL-GAS text pattern",
            )
        else:
            match = fee_pattern.match(combined)
            if match:
                quantity = text_to_digit(match.group("q1"))
                unit_price = parse_invoice_money(match.group("price"))
                amount = parse_invoice_money(match.group("amount"))
                if quantity == 1 and unit_price is not None and amount is not None and not close_money(unit_price, amount):
                    if amount > unit_price * 2:
                        amount = unit_price
                    else:
                        unit_price = amount
                description = " ".join(
                    part
                    for part in (match.group("prefix"), match.group("desc"))
                    if part
                ).strip()
                item = make_high_item(
                    match.group("code"),
                    description,
                    quantity,
                    unit_price,
                    match.group("unit"),
                    amount,
                    page_number,
                    "Aero ALL-GAS fee pattern",
                )
            else:
                match = relaxed_product_pattern.match(combined)
                if match:
                    unit_price = parse_invoice_money(match.group("price"))
                    amount = parse_invoice_money(match.group("amount"))
                    quantity = text_to_digit(match.group("qty"))
                    if (
                        quantity is None
                        or (
                            unit_price is not None
                            and amount is not None
                            and not close_money(quantity * unit_price, amount)
                        )
                    ):
                        if unit_price and amount:
                            calculated = round(amount / unit_price)
                            if close_money(calculated * unit_price, amount):
                                quantity = float(calculated)
                    item = make_high_item(
                        match.group("code"),
                        match.group("desc"),
                        quantity,
                        unit_price,
                        match.group("unit").replace("CYT", "CYL").replace("CYZ", "CYL"),
                        amount,
                        page_number,
                        "Aero ALL-GAS relaxed pattern",
                    )
                else:
                    match = repair_pattern.match(combined)
                    if match:
                        unit_price = parse_invoice_money(match.group("price"))
                        amount = parse_invoice_money(match.group("amount"))
                        quantity = text_to_digit(match.group("qty"))
                        if (
                            quantity is None
                            or (
                                unit_price is not None
                                and amount is not None
                                and not close_money(quantity * unit_price, amount)
                            )
                        ):
                            if unit_price and amount:
                                calculated = round(amount / unit_price)
                                if close_money(calculated * unit_price, amount):
                                    quantity = float(calculated)
                        item = make_high_item(
                            match.group("code"),
                            match.group("desc"),
                            quantity,
                            unit_price,
                            match.group("unit").replace("EBA", "EA"),
                            amount,
                            page_number,
                            "Aero ALL-GAS repair pattern",
                        )
                    else:
                        match = legacy_product_pattern.search(combined)
                        if not match:
                            continue
                        unit_price = parse_invoice_money(match.group("price"))
                        amount = parse_invoice_money(match.group("amount"))
                        quantity = billed_quantity_from_candidates(
                            [text_to_digit(match.group("q1")), text_to_digit(match.group("q2"))],
                            unit_price,
                            amount,
                        )
                        if quantity is not None and unit_price is not None and amount is not None and not close_money(quantity * unit_price, amount):
                            calculated = round(amount / unit_price) if unit_price else None
                            if calculated is not None and close_money(calculated * unit_price, amount):
                                quantity = float(calculated)
                        item = make_high_item(
                            match.group("code"),
                            match.group("desc"),
                            quantity,
                            unit_price,
                            "CYL",
                            amount,
                            page_number,
                            "Aero ALL-GAS legacy pattern",
                        )

        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def extract_airgenics_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    known_co2_prices = {
        "50": 17.00,
        "20": 8.50,
        "10": 5.00,
        "5": 4.00,
        "2.5": 2.00,
    }
    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace("|", " ").replace("]", " ")
        line = re.sub(r"CO[Z2][- ]?[S5][O0]{1,2}", "CO2-50", line, flags=re.IGNORECASE)
        if len(line) > 220 and re.search(r"\b5\s+1\s+1\s+1\b", line):
            continue
        upper = line.upper().replace("COZ", "CO2")
        money_tokens = re.findall(r"\d+[\.,]\d{2,4}", line)
        if "FUEL" in upper and "SURCHARGE" in upper and money_tokens:
            amount = parse_invoice_money(money_tokens[-1])
            item = make_high_item(
                "FUEL-DELIVERY",
                "FUEL AND DELIVERY SURCHARGE",
                1,
                amount,
                "EA",
                amount,
                page_number,
                "Airgenics surcharge pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
            continue
        if "CO2" not in upper:
            continue
        code_match = re.search(r"\b(?:CO2|COZ)-?\s*(2\.5|\d{1,3})\b", line, re.IGNORECASE)
        size_match = code_match or re.search(r"\b(50|20|10|5|2\.5)-?\s*LB\s+CO2\s+FILLED\b", upper, re.IGNORECASE)
        if not size_match:
            continue
        size = size_match.group(1)
        code = f"CO2-{size}"
        if len(money_tokens) >= 2:
            unit_price = parse_invoice_money(money_tokens[-2])
            amount = parse_invoice_money(money_tokens[-1])
        elif len(money_tokens) == 1 and size in known_co2_prices:
            unit_price = known_co2_prices[size]
            amount = parse_invoice_money(money_tokens[-1])
        else:
            continue
        qty_match = re.search(r"^\D*(\d{1,3})\D+(?:CO2|COZ)", line, re.IGNORECASE)
        quantity = parse_number(qty_match.group(1)) if qty_match else None
        if quantity is not None and unit_price and amount and not close_money(quantity * unit_price, amount):
            calculated = round(amount / unit_price)
            if close_money(calculated * unit_price, amount):
                quantity = float(calculated)
        if quantity is None and unit_price and amount:
            calculated = round(amount / unit_price)
            if close_money(calculated * unit_price, amount):
                quantity = float(calculated)
        description_prefix = (
            re.escape(code_match.group(0))
            if code_match
            else re.escape(size) + r"-?LB\s+CO2"
        )
        description_price = (
            re.escape(money_tokens[-2])
            if len(money_tokens) >= 2
            else re.escape(money_tokens[-1])
        )
        description_match = re.search(
            r"(?:" + description_prefix + r")\s+(?P<desc>.+?)\s+" + description_price,
            line,
            re.IGNORECASE,
        )
        description = description_match.group("desc") if description_match else f"{size}-LB CO2 FILLED"
        item = make_high_item(
            code,
            description,
            quantity,
            unit_price,
            "",
            amount,
            page_number,
            "Airgenics text pattern",
        )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def extract_alliance_energy_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    last_quantity: float | None = None
    main_pattern = re.compile(
        r"^(?P<date>\d{1,2}/\d{1,2}/\d{4})\s+"
        r"(?P<delivery>\d+)\s+Truck\s+"
        r"(?P<desc>Propane)\s+"
        r"(?P<qty>\d[\d,]*)\s+"
        r"(?P<unit>gal)\s+"
        r"(?P<price>\d+(?:\.\d{2,5})?)\s+=?\s*gal\s+"
        r"(?P<amount>\d[\d,]*\.\d{2})",
        re.IGNORECASE,
    )
    fee_pattern = re.compile(
        r"^(?P<desc>Federal PERC Fee|NJ PERC Fee)\s+"
        r"(?P<price>\d+(?:\.\d{2,5})?)\s+gal\s+"
        r"(?P<amount>\d[\d,]*\.\d{2})$",
        re.IGNORECASE,
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line)
        match = main_pattern.match(line)
        if match:
            quantity = parse_number(match.group("qty"))
            last_quantity = quantity
            items.append(
                make_high_item(
                    match.group("delivery"),
                    match.group("desc"),
                    quantity,
                    parse_invoice_money(match.group("price")),
                    match.group("unit"),
                    parse_invoice_money(match.group("amount")),
                    page_number,
                    "Alliance Energy text pattern",
                )
            )
            continue
        match = fee_pattern.match(line)
        if match:
            items.append(
                make_high_item(
                    "",
                    match.group("desc"),
                    last_quantity,
                    parse_invoice_money(match.group("price")),
                    "gal",
                    parse_invoice_money(match.group("amount")),
                    page_number,
                    "Alliance Energy fee pattern",
                )
            )
    return items


def extract_eastern_propane_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    money = r"\$?(?:\d[\d,]*|\.\d+)(?:\.\d{1,5})?"
    delivery_pattern = re.compile(
        rf"^(?P<date>\d{{1,2}}/\d{{1,2}}/\d{{2,4}})\.?\s+"
        rf"(?P<desc>.+?)\s+"
        rf"(?P<qty>\d[\d,]*(?:\.\d+)?)\s+"
        rf"(?P<price>{money})\s+"
        rf"(?P<amount>{money})$",
        re.IGNORECASE,
    )
    meter_pattern = re.compile(
        rf"^(?P<date>\d{{1,2}}/\d{{1,2}}/\d{{2,4}})\.?\s+"
        rf"(?P<qty>\d[\d,]*(?:\.\d+)?)\s+"
        rf"(?P<desc>Propane|Meter Maintenance Fee)\s+"
        rf"(?P<ticket>[A-Z0-9-]+)\s+"
        rf"(?P<price>{money})\s+"
        rf"(?P<amount>{money})$",
        re.IGNORECASE,
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace("|", " ")
        line = re.sub(r"\s+", " ", line).strip(" .|:;")
        if not line or re.search(r"Sales Tax|Tax Exempt|Amount Due|Delivery History|Other Ways", line, re.IGNORECASE):
            continue
        match = meter_pattern.match(line)
        item_number = ""
        unit = "GAL"
        if match:
            item_number = match.group("ticket")
        else:
            match = delivery_pattern.match(line)
            if not match:
                continue
            desc_head = match.group("desc").strip()
            if re.fullmatch(r"Delivery Charge", desc_head, re.IGNORECASE):
                amount = parse_invoice_money(match.group("amount"))
                if amount in (None, 0):
                    continue
            if not re.search(r"Propane|Hazardous|Compliance|Fuel Recovery|Delivery Charge|Meter Maintenance", desc_head, re.IGNORECASE):
                continue

        quantity = parse_invoice_money(match.group("qty"))
        unit_price = parse_invoice_money(match.group("price"))
        amount = reconcile_amount(quantity, unit_price, parse_invoice_money(match.group("amount")))
        description = re.sub(r"\s+", " ", match.group("desc")).strip(" -")
        if not close_money((quantity or 0) * (unit_price or 0), amount):
            confidence = "Medium"
        else:
            confidence = "High"
        item = LineItem(
            item_number=item_number,
            description=clean_ocr_text(description),
            quantity=quantity,
            unit_price=unit_price,
            unit=unit,
            line_total=amount,
            page=page_number,
            confidence=confidence,
            extraction_method="Eastern Propane text pattern",
        )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def extract_idm_medical_gas_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    pattern = re.compile(
        r"^(?P<qty>\d+(?:\.\d+)?)\s+"
        r"(?P<code>[A-Z0-9/-]+)\s+"
        r"(?P<desc>.+?)\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d[\d,.]*[\.,]\d{2})$",
        re.IGNORECASE,
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace("|", " ")
        line = re.sub(r"\s+", " ", line).strip(" .|")
        match = pattern.match(line)
        if not match or not re.search(r"REFILL|OXYGEN|NITROGEN|CYL", match.group("desc"), re.IGNORECASE):
            continue
        quantity = parse_invoice_money(match.group("qty"))
        unit_price = parse_invoice_money(match.group("price"))
        amount = reconcile_amount(quantity, unit_price, parse_invoice_money(match.group("amount")))
        item = make_high_item(
            match.group("code"),
            match.group("desc"),
            quantity,
            unit_price,
            "CYL",
            amount,
            page_number,
            "IDM Medical Gas text pattern",
        )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def extract_c_three_logistics_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    freight_pattern = re.compile(
        r"^(?P<order>\d{5,})\s+(?P<bol>\d+)\s+"
        r"(?P<date>\d{1,2}/\d{1,2}/\d{4})\s+"
        r"(?P<units>\d[\d,]*(?:\.\d+)?)\s+"
        r"(?P<rate>\.?\d+(?:\.\d{2,5})?)\s+"
        r"(?P<amount>\d[\d,]*(?:\.\d{2})?)$",
        re.IGNORECASE,
    )
    fee_pattern = re.compile(
        r"^(?P<desc>Fuel\s+Sur\s+Charge|Tolls?)\s+"
        r"(?P<units>\d[\d,]*(?:\.\d+)?)\s+"
        r"(?P<rate>\d+(?:\.\d{2,5})?)\s+"
        r"(?P<amount>\d[\d,]*(?:\.\d{2})?)$",
        re.IGNORECASE,
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace("|", " ")
        line = re.sub(r"\s+", " ", line).strip(" .|")
        match = freight_pattern.match(line)
        if match:
            quantity = parse_invoice_money(match.group("units"))
            unit_price = parse_invoice_money(match.group("rate"))
            amount = reconcile_amount(quantity, unit_price, parse_invoice_money(match.group("amount")))
            item = make_high_item(
                match.group("order"),
                f"Freight delivery BOL {match.group('bol')}",
                quantity,
                unit_price,
                "EA",
                amount,
                page_number,
                "C Three Logistics text pattern",
            )
        else:
            match = fee_pattern.match(line)
            if not match:
                continue
            quantity = parse_invoice_money(match.group("units"))
            unit_price = parse_invoice_money(match.group("rate"))
            amount = parse_invoice_money(match.group("amount"))
            item = make_high_item(
                "",
                match.group("desc"),
                quantity,
                unit_price,
                "EA",
                amount,
                page_number,
                "C Three Logistics fee pattern",
            )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def extract_american_compressed_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    embedded_pattern = re.compile(
        r"(?P<code>\d{6})\s+"
        r"(?P<ticket>[A-Z]{1,3}\d{2,5})\s+"
        r"(?P<qty>\d+|[!Il])\s+"
        r"(?P<body>.+?)\s+"
        r"(?:(?P<unit>EA|CYL|MIS|REFIL|REPFIL|RE-FIL|REFILL)\s+)?"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})"
        r"(?=\s+\d{6}\s+[A-Z]{1,3}\d{2,5}\s+|\s+SUB\s*TOTAL|\s*$)",
        re.IGNORECASE,
    )
    pattern = re.compile(
        r"^(?P<code>\d{6})\s+"
        r"(?P<ticket>[A-Z0-9]+)\s+"
        r"(?P<qty>\d+|[!Il])\s+"
        r"(?P<body>.+?)\s+"
        r"(?P<unit>EA|CYL|MIS|REFIL|REPFIL|RE-FIL|REFILL)\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})$",
        re.IGNORECASE,
    )
    no_unit_pattern = re.compile(
        r"^(?P<code>\d{6})\s+"
        r"(?P<ticket>[A-Z0-9]+)\s+"
        r"(?P<qty>\d+|[!Il])\s+"
        r"(?P<body>.+?)\s+"
        r"(?P<price>\d+[\.,]\d{2,4})\s+"
        r"(?P<amount>\d+[\.,]\d{2,4})$",
        re.IGNORECASE,
    )
    for raw_line in lines:
        line = clean_ocr_text(raw_line).replace("|", " ").replace("—", " ")
        matches = list(embedded_pattern.finditer(line))
        if not matches:
            match = pattern.match(line)
            if match:
                matches = [match]
            else:
                match = no_unit_pattern.match(line)
                if match:
                    matches = [match]
        for match in matches:
            unit = ""
            if "unit" in match.groupdict() and match.group("unit"):
                unit = match.group("unit").upper()
                if unit in {"REFIL", "REPFIL", "RE-FIL", "REFILL"}:
                    unit = "EA"
            else:
                unit = "EA"
            quantity = text_to_digit(match.group("qty"))
            unit_price = parse_invoice_money(match.group("price"))
            amount = parse_invoice_money(match.group("amount"))
            if quantity is not None and unit_price is not None and amount is not None and not close_money(quantity * unit_price, amount):
                for divisor in (10, 100, 1000):
                    shifted = unit_price / divisor
                    if close_money(quantity * shifted, amount):
                        unit_price = shifted
                        break
                else:
                    if quantity:
                        unit_price = round(amount / quantity, 4)
            description = re.sub(r"\s+", " ", match.group("body")).strip()
            item = make_high_item(
                match.group("code"),
                description,
                quantity,
                unit_price,
                unit,
                amount,
                page_number,
                "American Compressed text pattern",
            )
            key = (item.item_number, item.description, item.quantity, item.line_total)
            if key not in seen:
                seen.add(key)
                items.append(item)
        if matches:
            continue
    return items


def extract_atlas_items(lines: list[str], page_number: int) -> list[LineItem]:
    items: list[LineItem] = []
    seen: set[tuple[str, str, float | None, float | None]] = set()
    for index, raw_line in enumerate(lines):
        line = clean_ocr_text(raw_line).replace("|", " ")
        if not re.search(r"\b(?:CYL|EA)\b", line, re.IGNORECASE):
            continue
        money_tokens = re.findall(r"\d+[\.,]\d{2,4}", line)
        if len(money_tokens) < 2:
            continue
        unit_match = re.search(r"\b(CYL|EA)\b", line, re.IGNORECASE)
        if not unit_match:
            continue
        quantity_match = re.search(r"(\d+(?:[\.,]\d+)?)\W*\s+" + unit_match.group(1), line, re.IGNORECASE)
        quantity = parse_invoice_money(quantity_match.group(1)) if quantity_match else None
        unit_price = parse_invoice_money(money_tokens[-2])
        amount = parse_invoice_money(money_tokens[-1])
        if quantity is not None and unit_price is not None and amount is not None:
            calculated = amount / unit_price if unit_price else None
            if calculated is not None and close_money(round(calculated) * unit_price, amount):
                quantity = float(round(calculated))

        code_match = re.search(
            r"[}\s](?P<code>[A-Z0-9]{2,}(?:\s+[A-Z0-9]{2,})?)\s+"
            r"(?:\d+[lI|]?\s*[0O]?\s+)?"
            + re.escape(money_tokens[-2]),
            line,
            re.IGNORECASE,
        )
        item_number = code_match.group("code") if code_match else ""
        desc_parts: list[str] = []
        for following in lines[index + 1:index + 4]:
            following_clean = clean_ocr_text(following).strip(" |")
            if not following_clean or "CUSTOMER OWNED" in following_clean.upper():
                break
            if re.search(r"SUBTOTAL|TOTAL|CHEMTREC|UN\d{4}", following_clean, re.IGNORECASE):
                break
            desc_parts.append(following_clean)
        description = " ".join(desc_parts) or item_number or "Atlas invoice item"
        item = make_high_item(
            item_number,
            description,
            quantity,
            unit_price,
            unit_match.group(1),
            amount,
            page_number,
            "Atlas text pattern",
        )
        key = (item.item_number, item.description, item.quantity, item.line_total)
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def extract_vendor_specific_items(lines: list[str], page_number: int) -> list[LineItem]:
    joined = "\n".join(lines)
    if re.search(r"\b(?:Aero\s+)?ALL-?GAS\b|The Aero ALL-GAS Co\.|www\.?aallgas\.com|278-AERO", joined, re.IGNORECASE):
        return extract_aero_all_gas_items(lines, page_number)
    if re.search(r"\bATLAS\b|[ae]tlaswe?id[sd]\.[a-z]{2,4}|BROOK\s*ROA[DN]|LAKEWOOD", joined, re.IGNORECASE):
        return extract_atlas_items(lines, page_number)
    if re.search(r"\bEastern Propane\b|easternpropane\.com|Superior\s+Plus.*Propane", joined, re.IGNORECASE | re.DOTALL):
        return extract_eastern_propane_items(lines, page_number)
    if re.search(r"\bI\.?\s*D\.?\s*M\.?\s+MEDICAL\s+GAS\b|MEDICAL GAS REFILLS", joined, re.IGNORECASE):
        return extract_idm_medical_gas_items(lines, page_number)
    if re.search(r"\bC\s*Three\s+Logistics\s+LLC\b|Logistics,\s*LLC.*PO Box 914", joined, re.IGNORECASE | re.DOTALL):
        return extract_c_three_logistics_items(lines, page_number)
    if re.search(r"\bAIRGENICS\b|Premium Helium|Manufacturers Place|Caroline@Airgenics", joined, re.IGNORECASE):
        return extract_airgenics_items(lines, page_number)
    if re.search(r"\bAlliance Energy Services\b|\bALLIANCE\b.*SALES INVOICE", joined, re.IGNORECASE):
        return extract_alliance_energy_items(lines, page_number)
    if re.search(r"\bAMERICAN COMPRESSED GASES\b", joined, re.IGNORECASE):
        return extract_american_compressed_items(lines, page_number)
    if re.search(r"\bWESTERN\b|INTERNATIONAL GAS & CYLINDERS", joined, re.IGNORECASE):
        return extract_western_items(lines, page_number)
    if re.search(r"\bMESSER\b|Gases for Life", joined, re.IGNORECASE):
        return extract_messer_items(lines, page_number)
    if re.search(r"\bLinde Inc\b|PO Box 417518|PRODUCT DESCRIPTION, ORDER", joined, re.IGNORECASE):
        return extract_linde_bulk_items(lines, page_number)
    if re.search(r"Linde\d?\s+Gas\d?\s*&\d?\s*Equipment\d?|Linde Gas\s*&\s*Equipment|ITEM NUMBER ITEM DESCRIPTION", joined, re.IGNORECASE):
        return extract_linde_gas_equipment_items(lines, page_number)
    return []


def ocr_page(pdf_path: Path, page_number: int, dpi: int) -> str:
    if not shutil.which("tesseract"):
        raise RuntimeError(
            "This PDF needs OCR, but Tesseract is not installed or is not on PATH."
        )
    images = convert_from_path(
        str(pdf_path),
        dpi=dpi,
        first_page=page_number,
        last_page=page_number,
        fmt="png",
        thread_count=1,
    )
    if not images:
        return ""
    with tempfile.TemporaryDirectory(prefix="invoice_ocr_") as temp_dir:
        image_path = Path(temp_dir) / "page.png"
        images[0].save(image_path)
        completed = subprocess.run(
            [
                "tesseract", str(image_path), "stdout",
                "--dpi", str(dpi), "--psm", "6", "-l", "eng",
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        return completed.stdout


def process_pdf(pdf_path: Path, dpi: int = 250) -> Invoice:
    invoice = Invoice(source_file=pdf_path.name)
    table_items: list[LineItem] = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                usable_chars = len(re.sub(r"\s+", "", text))
                if usable_chars < 40:
                    text = ocr_page(pdf_path, page_number, dpi)
                    invoice.notes.append(f"Page {page_number}: OCR used")
                else:
                    page_items = extract_table_items(page, page_number)
                    table_items.extend(page_items)
                invoice.raw_pages.append(text)
    except Exception as exc:
        raise RuntimeError(f"Could not process {pdf_path.name}: {exc}") from exc

    all_text = "\n".join(invoice.raw_pages)
    lines = [clean_ocr_text(line) for line in all_text.splitlines() if clean_cell(line)]
    invoice.vendor_name = extract_vendor(lines)
    invoice.invoice_date = extract_date(lines)
    invoice.invoice_number = extract_invoice_number(lines)
    if not invoice.invoice_number:
        invoice.invoice_number = item_number_from_filename(pdf_path)
    invoice.ship_to_address = extract_ship_to(lines)

    invoice.items = table_items
    if not invoice.items:
        for page_number, page_text in enumerate(invoice.raw_pages, start=1):
            page_lines = [
                clean_ocr_text(line) for line in page_text.splitlines() if clean_cell(line)
            ]
            invoice.items.extend(extract_vendor_specific_items(page_lines, page_number))
    if not invoice.items:
        for page_number, page_text in enumerate(invoice.raw_pages, start=1):
            page_lines = [
                clean_ocr_text(line) for line in page_text.splitlines() if clean_cell(line)
            ]
            invoice.items.extend(extract_text_items(page_lines, page_number))

    if not invoice.vendor_name:
        invoice.notes.append("Vendor name not found")
    if not invoice.invoice_date:
        invoice.notes.append("Invoice date not found")
    if not invoice.ship_to_address:
        invoice.notes.append("Ship-to address not found")
    if not invoice.items:
        invoice.notes.append("No line items found")
    return invoice


def find_pdfs(input_path: Path) -> list[Path]:
    if input_path.is_file():
        if input_path.suffix.lower() != ".pdf":
            raise ValueError("The input file must be a PDF.")
        return [input_path]
    if not input_path.is_dir():
        raise ValueError(f"Input path does not exist: {input_path}")
    return sorted(
        path for path in input_path.rglob("*")
        if path.is_file() and path.suffix.lower() == ".pdf"
    )


def style_sheet(sheet, widths: dict[str, int]) -> None:
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def write_workbook(invoices: list[Invoice], output_path: Path) -> None:
    workbook = Workbook()
    items_sheet = workbook.active
    items_sheet.title = "Line Items"
    item_headers = [
        "Source PDF", "Page", "Vendor Name", "Invoice Date", "Invoice Number",
        "Ship To Address", "Item Number", "Description", "Quantity",
        "Unit Price", "Unit", "Line Total", "Confidence", "Extraction Method",
    ]
    items_sheet.append(item_headers)

    invoice_sheet = workbook.create_sheet("Invoice Summary")
    invoice_sheet.append(
        [
            "Source PDF", "Vendor Name", "Invoice Date", "Invoice Number",
            "Ship To Address", "Line Item Count", "Notes",
        ]
    )

    review_sheet = workbook.create_sheet("Needs Review")
    review_sheet.append(
        [
            "Source PDF", "Page", "Problem", "Vendor Name", "Invoice Date",
            "Item Number", "Description", "Quantity", "Unit Price", "Unit",
        ]
    )

    raw_sheet = workbook.create_sheet("Raw Text")
    raw_sheet.append(["Source PDF", "Page", "Extracted Text"])

    for invoice in invoices:
        invoice_sheet.append(
            [
                invoice.source_file, invoice.vendor_name, invoice.invoice_date,
                invoice.invoice_number, invoice.ship_to_address, len(invoice.items),
                "; ".join(invoice.notes),
            ]
        )
        for page_number, text in enumerate(invoice.raw_pages, start=1):
            raw_sheet.append([invoice.source_file, page_number, text])

        if not invoice.items:
            review_sheet.append(
                [
                    invoice.source_file, "", "No line items found",
                    invoice.vendor_name, invoice.invoice_date, "", "", "", "", "",
                ]
            )

        for item in invoice.items:
            items_sheet.append(
                [
                    invoice.source_file, item.page, invoice.vendor_name,
                    invoice.invoice_date, invoice.invoice_number,
                    invoice.ship_to_address, item.item_number, item.description,
                    item.quantity, item.unit_price, item.unit, item.line_total,
                    item.confidence, item.extraction_method,
                ]
            )
            if item.confidence == "Review":
                review_sheet.append(
                    [
                        invoice.source_file, item.page, "Check extracted line item",
                        invoice.vendor_name, invoice.invoice_date, item.item_number,
                        item.description, item.quantity, item.unit_price, item.unit,
                    ]
                )

        for note in invoice.notes:
            if "not found" in note.lower():
                review_sheet.append(
                    [
                        invoice.source_file, "", note, invoice.vendor_name,
                        invoice.invoice_date, "", "", "", "", "",
                    ]
                )

    currency_format = '$#,##0.00'
    for row in range(2, items_sheet.max_row + 1):
        items_sheet.cell(row, 10).number_format = currency_format
        items_sheet.cell(row, 12).number_format = currency_format
    for row in range(2, review_sheet.max_row + 1):
        review_sheet.cell(row, 9).number_format = currency_format

    style_sheet(
        items_sheet,
        {
            "A": 28, "B": 8, "C": 25, "D": 14, "E": 16, "F": 42,
            "G": 18, "H": 45, "I": 12, "J": 14, "K": 10, "L": 14,
            "M": 12, "N": 18,
        },
    )
    style_sheet(
        invoice_sheet,
        {"A": 28, "B": 25, "C": 14, "D": 16, "E": 45, "F": 15, "G": 45},
    )
    style_sheet(
        review_sheet,
        {
            "A": 28, "B": 8, "C": 30, "D": 25, "E": 14,
            "F": 18, "G": 45, "H": 12, "I": 14, "J": 10,
        },
    )
    style_sheet(raw_sheet, {"A": 28, "B": 8, "C": 120})
    raw_sheet.sheet_view.showGridLines = False
    raw_sheet.column_dimensions["C"].width = 120
    for row in range(2, raw_sheet.max_row + 1):
        raw_sheet.row_dimensions[row].height = 90

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)


# Standalone runner starts here.
REQUIRED_RULES_VERSION = "2026-06-24.3"
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
    ".txt",
    ".docx",
}


@dataclass
class OcrResult:
    text: str
    confidence: float | None = None
    method: str = ""
    cache_hit: bool = False


@dataclass
class SourceDiagnostics:
    source_file: str
    source_type: str
    pages: int = 0
    ocr_pages: int = 0
    avg_ocr_confidence: float | None = None
    notes: list[str] = field(default_factory=list)


def parse_tesseract_tsv(tsv_text: str) -> tuple[str, float | None]:
    reader = csv.DictReader(tsv_text.splitlines(), delimiter="\t")
    lines: dict[tuple[int, int, int], list[str]] = {}
    confidences: list[float] = []

    for row in reader:
        text = clean_cell(row.get("text"))
        if not text:
            continue
        try:
            conf = float(row.get("conf", "-1"))
        except ValueError:
            conf = -1
        if conf >= 0:
            confidences.append(conf)

        try:
            key = (
                int(row.get("page_num", 0)),
                int(row.get("block_num", 0)),
                int(row.get("line_num", 0)),
            )
        except ValueError:
            key = (0, 0, 0)
        lines.setdefault(key, []).append(text)

    text_lines = [" ".join(parts) for _, parts in sorted(lines.items())]
    confidence = mean(confidences) if confidences else None
    return "\n".join(text_lines), confidence


def preprocess_image(image: Image.Image, scale: int, threshold: bool) -> Image.Image:
    processed = image.convert("L")
    processed = ImageOps.autocontrast(processed)
    if scale > 1:
        processed = processed.resize(
            (processed.width * scale, processed.height * scale),
            Image.Resampling.LANCZOS,
        )
    processed = processed.filter(ImageFilter.SHARPEN)
    if threshold:
        processed = processed.point(lambda pixel: 255 if pixel > 180 else 0)
    return processed


def run_tesseract(image: Image.Image, psm: int, scale: int, threshold: bool) -> OcrResult:
    if not shutil.which("tesseract"):
        raise RuntimeError("Tesseract is not installed or not on PATH.")

    with tempfile.TemporaryDirectory(prefix="invoice_ocr_") as temp_dir:
        image_path = Path(temp_dir) / "page.png"
        preprocess_image(image, scale=scale, threshold=threshold).save(image_path)
        completed = subprocess.run(
            [
                "tesseract",
                str(image_path),
                "stdout",
                "--psm",
                str(psm),
                "-l",
                "eng",
                "tsv",
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=60,
        )
    text, confidence = parse_tesseract_tsv(completed.stdout)
    return OcrResult(
        text=text,
        confidence=confidence,
        method=f"tesseract psm={psm} scale={scale} threshold={threshold}",
    )


def ocr_attempts(ocr_mode: str) -> list[tuple[int, int, bool]]:
    if ocr_mode == "fast":
        return [(6, 2, False)]
    if ocr_mode == "deep":
        return [(6, 2, False), (4, 2, False), (6, 2, True), (6, 3, False), (11, 2, False)]
    return [(6, 2, False), (4, 2, False), (6, 2, True)]


def ocr_score(result: OcrResult) -> tuple[float, int]:
    confidence = result.confidence if result.confidence is not None else 0
    useful_chars = len(re.sub(r"\s+", "", result.text))
    return confidence, useful_chars


def good_enough_ocr(result: OcrResult) -> bool:
    confidence, useful_chars = ocr_score(result)
    return confidence >= 68 and useful_chars >= 80


def ocr_image(image: Image.Image, ocr_mode: str = "balanced") -> OcrResult:
    candidates: list[OcrResult] = []

    for attempt_index, (psm, scale, threshold) in enumerate(ocr_attempts(ocr_mode)):
        try:
            result = run_tesseract(image, psm=psm, scale=scale, threshold=threshold)
        except Exception:
            continue
        if result.text.strip():
            candidates.append(result)
        if attempt_index == 0 and ocr_mode in {"balanced", "deep"} and good_enough_ocr(result):
            return result

    if not candidates:
        return OcrResult(text="", confidence=None, method="tesseract failed")

    return max(candidates, key=ocr_score)


def source_signature(path: Path) -> str:
    stat = path.stat()
    return f"{path.resolve()}|{stat.st_size}|{stat.st_mtime_ns}"


def cache_key(path: Path, page_number: int, dpi: int, ocr_mode: str) -> str:
    digest = hashlib.sha256(
        f"{source_signature(path)}|page={page_number}|dpi={dpi}|ocr={ocr_mode}|v=3".encode("utf-8")
    ).hexdigest()
    return digest


def read_ocr_cache(cache_dir: Path | None, key: str) -> OcrResult | None:
    if cache_dir is None:
        return None
    cache_path = cache_dir / f"{key}.json"
    if not cache_path.exists():
        return None
    try:
        data = json.loads(cache_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return OcrResult(
        text=str(data.get("text", "")),
        confidence=data.get("confidence"),
        method=str(data.get("method", "cached OCR")),
        cache_hit=True,
    )


def write_ocr_cache(cache_dir: Path | None, key: str, result: OcrResult) -> None:
    if cache_dir is None:
        return
    try:
        cache_dir.mkdir(parents=True, exist_ok=True)
        cache_path = cache_dir / f"{key}.json"
        cache_path.write_text(
            json.dumps(
                {
                    "text": result.text,
                    "confidence": result.confidence,
                    "method": result.method,
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
    except Exception:
        return


def find_files(input_path: Path) -> list[Path]:
    if input_path.is_file():
        if input_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {input_path.suffix}")
        return [input_path]
    if not input_path.is_dir():
        raise ValueError(f"Input path does not exist: {input_path}")
    return sorted(
        path
        for path in input_path.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def lines_from_pages(pages: Iterable[str]) -> list[str]:
    all_text = "\n".join(pages)
    return [clean_ocr_text(line) for line in all_text.splitlines() if clean_cell(line)]


def vendor_from_source_name(source_file: str) -> str:
    upper = source_file.upper()
    if "AERO ALL-GAS" in upper or "AER ALL-GAS" in upper or "ALL GAS" in upper:
        return "The Aero ALL-GAS Co."
    if "AIRGENICS" in upper:
        return "Airgenics LLC"
    if "ALLIANCE" in upper or "ALLIANCE NERGY" in upper:
        return "Alliance Energy Services LLC"
    if "AMERICAN COMPRESSED" in upper:
        return "American Compressed Gases Inc."
    if "ATLAS WELDING" in upper or upper.startswith("ATLAS "):
        return "Atlas Welding Supply Co."
    if "EASTERN PROPANE" in upper:
        return "Eastern Propane"
    if "IDM MEDICAL" in upper or "I.D.M" in upper:
        return "I.D.M. Medical Gas Co."
    if "C THREE" in upper:
        return "C Three Logistics LLC"
    if "MAT AMERICA" in upper:
        return "MAT America Inc."
    if "MESSER" in upper:
        return "Messer LLC"
    if "WESTERN" in upper:
        return "Western Intl Gas & Cylinders, Inc."
    if "LINDE GAS" in upper or "LINDE INC" in upper or "LINDE" in upper:
        if "EQUIP" in upper or "EQQUIP" in upper or "EQUIPMENT" in upper:
            return "Linde Gas & Equipment Inc."
        return "Linde Inc."
    return ""


def ship_to_from_source_and_text(source_file: str, pages: list[str]) -> str:
    upper_name = source_file.upper()
    joined = "\n".join(pages)
    joined_upper = joined.upper()
    if "AERO ALL-GAS" in upper_name or "AER ALL-GAS" in upper_name or "ALL GAS" in upper_name:
        return "AWISCO, 465 West Main Street, STAMFORD CT 06902"
    if "EASTERN PROPANE" in upper_name:
        service_match = re.search(
            r"Service\s+Address\s*:?\s*(?P<addr>.+?)(?:\n|Date\s+(?:Item|Gallons)|\d{1,2}/\d{1,2}/\d{2,4})",
            joined,
            re.IGNORECASE | re.DOTALL,
        )
        if service_match:
            address = clean_ocr_text(service_match.group("addr"))
            address = re.sub(r"\b(?:Date|Item|Description|Quantity|Unit Price|Amount)\b.*$", "", address, flags=re.IGNORECASE).strip(" :|-")
            address = re.sub(
                r"(?<=[A-Z])(?=\d{1,5}\s)|(?<=[a-z])(?=[A-Z]{2,}\b)|(?<=DRIVE)(?=WEST)|(?<=206)(?=SANDYSTON|BRANCHVILLE)",
                ", ",
                address,
            )
            if len(address) >= 10:
                return address
    if "C THREE" in upper_name and "WOODBRIDGE" in joined_upper:
        return "AWISCO, 475 US Highway 9S, WOODBRIDGE, NJ 07095"
    if "IDM MEDICAL" in upper_name or "I.D.M" in upper_name:
        if "WEST ORANGE" in joined_upper:
            return "GAS ARC COMPANY INC, 24 LAKESIDE AND STANDISH AVE, WEST ORANGE, NJ 07052"
    if "AIRGENICS" in upper_name:
        if "ASCO" in joined_upper or "475" in joined_upper or "PRINCETON" in joined_upper:
            return "AWISCO / ASCO, 475 U.S. ROUTE 9 SOUTH, PRINCETON, NJ 08550"
    if "ATLAS WELDING" in upper_name:
        if "TOMS" in joined_upper or "1889" in joined_upper:
            return "AWISCO, 1889 ROUTE 9, SUITE 116, TOMS RIVER NJ 08755"
    if "LINDE GAS" in upper_name and "EQUIP" in upper_name:
        if "WOODBRIDGE" in joined_upper and ("475RTE9S" in joined_upper or "475 RTE 9" in joined_upper):
            return "AWISCO ASCO, 475 RTE 9 S, WOODBRIDGE NJ 07095"
        if "WEST ORANGE" in joined_upper and "24 LAKESIDE" in joined_upper:
            return "AWISCO CORP, 24 LAKESIDE AVE, WEST ORANGE NJ 07052"
    if "AMERICAN COMPRESSED" in upper_name:
        if "2660 PARK" in joined_upper and "BRONX" in joined_upper:
            if "10451-6015" in joined_upper:
                return "AWISCO BRONX, 2660 PARK AVENUE, BRONX NY 10451-6015"
            return "AWISCO, 2660 PARK AVE, BRONX NY"
    if "ALLIANCE" in upper_name or "ALLIANCE NERGY" in upper_name:
        if "WOODBRIDGE" in joined_upper:
            return "AWISCO NJ LLC, 475 Route 9S, Woodbridge, NJ 07095"
    return ""


def build_invoice_from_pages(
    source_file: str,
    pages: list[str],
    notes: list[str],
    table_items: list[LineItem] | None = None,
) -> Invoice:
    invoice = Invoice(source_file=source_file, raw_pages=pages, notes=list(notes))
    lines = lines_from_pages(pages)

    invoice.vendor_name = extract_vendor(lines)
    source_vendor = vendor_from_source_name(source_file)
    if source_vendor:
        invoice.vendor_name = source_vendor
    invoice.invoice_date = extract_date(lines)
    invoice.invoice_number = extract_invoice_number(lines)
    filename_invoice_number = item_number_from_filename(Path(source_file))
    if (
        not invoice.invoice_number
        or invoice.invoice_number.upper() in {"AND", "ANDDATE", "DATE", "NUMBER", "INVOICE", "POBOX", "BOX"}
        or (
            filename_invoice_number
            and invoice.invoice_number
            and re.fullmatch(r"[A-Z]+", invoice.invoice_number.upper())
            and re.search(r"\d", filename_invoice_number)
        )
        or (
            filename_invoice_number
            and invoice.invoice_number
            and filename_invoice_number not in invoice.invoice_number
            and invoice.invoice_number not in filename_invoice_number
            and len(invoice.invoice_number) < 5
        )
    ):
        invoice.invoice_number = filename_invoice_number
    invoice.ship_to_address = extract_ship_to(lines)
    source_ship_to = ship_to_from_source_and_text(source_file, pages)
    if source_ship_to and (
        not invoice.ship_to_address
        or len(invoice.ship_to_address) > 180
        or "ACCOUNT:" in invoice.ship_to_address.upper()
    ):
        invoice.ship_to_address = source_ship_to

    invoice.items = list(table_items or [])
    vendor_items: list[LineItem] = []
    for page_number, page_text in enumerate(pages, start=1):
        page_lines = [
            clean_ocr_text(line) for line in page_text.splitlines() if clean_cell(line)
        ]
        if source_vendor:
            page_lines = [source_vendor, *page_lines]
        vendor_items.extend(extract_vendor_specific_items(page_lines, page_number))
    if vendor_items and (
        not invoice.items
        or item_list_score(vendor_items) > item_list_score(invoice.items)
    ):
        invoice.items = vendor_items
    if not invoice.items:
        for page_number, page_text in enumerate(pages, start=1):
            page_lines = [
                clean_ocr_text(line) for line in page_text.splitlines() if clean_cell(line)
            ]
            invoice.items.extend(extract_text_items(page_lines, page_number))

    if not invoice.vendor_name:
        invoice.notes.append("Vendor name not found")
    if not invoice.invoice_date:
        invoice.notes.append("Invoice date not found")
    if not invoice.invoice_number:
        invoice.notes.append("Invoice number not found")
    if not invoice.ship_to_address:
        invoice.notes.append("Ship-to address not found")
    if not invoice.items:
        invoice.notes.append("No line items found")

    return invoice


def item_list_score(items: list[LineItem]) -> tuple[int, int, int, int]:
    high = sum(1 for item in items if item.confidence == "High")
    complete = sum(
        1
        for item in items
        if item.description and item.quantity is not None and (item.unit_price is not None or item.line_total is not None)
    )
    priced = sum(1 for item in items if item.unit_price is not None or item.line_total is not None)
    return (high, complete, priced, len(items))


def invoice_has_core_data(source_file: str, pages: list[str], table_items: list[LineItem]) -> bool:
    if not pages:
        return False
    invoice = build_invoice_from_pages(source_file, pages, [], table_items)
    return bool(
        invoice.vendor_name
        and invoice.invoice_date
        and invoice.invoice_number
        and invoice.items
    )


def ocr_pdf_page(
    pdf_path: Path,
    page_number: int,
    dpi: int,
    ocr_mode: str,
    cache_dir: Path | None,
) -> OcrResult:
    key = cache_key(pdf_path, page_number, dpi, ocr_mode)
    cached = read_ocr_cache(cache_dir, key)
    if cached:
        return cached

    images = convert_from_path(
        str(pdf_path),
        dpi=dpi,
        first_page=page_number,
        last_page=page_number,
        fmt="png",
        thread_count=1,
    )
    result = ocr_image(images[0], ocr_mode=ocr_mode) if images else OcrResult(text="")
    write_ocr_cache(cache_dir, key, result)
    return result


def process_pdf(
    pdf_path: Path,
    dpi: int,
    ocr_mode: str,
    cache_dir: Path | None,
    all_pages: bool,
) -> tuple[Invoice, SourceDiagnostics]:
    pages: list[str] = []
    table_items: list[LineItem] = []
    notes: list[str] = []
    ocr_confidences: list[float] = []
    page_count = 0

    with pdfplumber.open(pdf_path) as pdf:
        page_count = len(pdf.pages)
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            usable_chars = len(re.sub(r"\s+", "", text))
            if usable_chars >= 80:
                pages.append(text)
                table_items.extend(extract_table_items(page, page_number))
                continue

            if page_number > 1 and not all_pages and invoice_has_core_data(pdf_path.name, pages, table_items):
                pages.append("[Skipped OCR: invoice fields and line items were already found on earlier pages. Use --all-pages to force OCR on attachments.]")
                notes.append(f"Page {page_number}: OCR skipped after invoice data was found")
                continue

            result = ocr_pdf_page(
                pdf_path,
                page_number=page_number,
                dpi=dpi,
                ocr_mode=ocr_mode,
                cache_dir=cache_dir,
            )
            pages.append(result.text)
            if result.confidence is not None:
                ocr_confidences.append(result.confidence)
            conf_text = (
                f"{result.confidence:.1f}" if result.confidence is not None else "unknown"
            )
            cache_text = " from cache" if result.cache_hit else ""
            notes.append(f"Page {page_number}: OCR used{cache_text}, confidence {conf_text}")
            if result.confidence is None or result.confidence < 55:
                notes.append(f"Page {page_number}: low OCR confidence; review scan/handwriting")

    invoice = build_invoice_from_pages(pdf_path.name, pages, notes, table_items)
    diagnostics = SourceDiagnostics(
        source_file=pdf_path.name,
        source_type="PDF",
        pages=page_count,
        ocr_pages=sum(1 for note in notes if "OCR used" in note),
        avg_ocr_confidence=mean(ocr_confidences) if ocr_confidences else None,
        notes=notes,
    )
    return invoice, diagnostics


def process_image(
    image_path: Path,
    ocr_mode: str,
    cache_dir: Path | None,
) -> tuple[Invoice, SourceDiagnostics]:
    key = cache_key(image_path, 1, 0, ocr_mode)
    result = read_ocr_cache(cache_dir, key)
    if result is None:
        with Image.open(image_path) as image:
            result = ocr_image(image, ocr_mode=ocr_mode)
        write_ocr_cache(cache_dir, key, result)
    notes = ["Image OCR used" + (" from cache" if result.cache_hit else "")]
    if result.confidence is not None:
        notes.append(f"OCR confidence {result.confidence:.1f}")
    if result.confidence is None or result.confidence < 55:
        notes.append("Low OCR confidence; review image/handwriting")
    invoice = build_invoice_from_pages(image_path.name, [result.text], notes)
    diagnostics = SourceDiagnostics(
        source_file=image_path.name,
        source_type="Image",
        pages=1,
        ocr_pages=1,
        avg_ocr_confidence=result.confidence,
        notes=notes,
    )
    return invoice, diagnostics


def process_txt(text_path: Path) -> tuple[Invoice, SourceDiagnostics]:
    text = text_path.read_text(encoding="utf-8", errors="replace")
    invoice = build_invoice_from_pages(text_path.name, [text], ["Text file read directly"])
    diagnostics = SourceDiagnostics(
        source_file=text_path.name,
        source_type="TXT",
        pages=1,
        notes=["Text file read directly"],
    )
    return invoice, diagnostics


def process_docx(docx_path: Path) -> tuple[Invoice, SourceDiagnostics]:
    document = Document(docx_path)
    blocks: list[str] = []
    blocks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(clean_cell(cell.text) for cell in row.cells))
    text = "\n".join(blocks)
    invoice = build_invoice_from_pages(docx_path.name, [text], ["DOCX text read directly"])
    diagnostics = SourceDiagnostics(
        source_file=docx_path.name,
        source_type="DOCX",
        pages=1,
        notes=["DOCX text read directly"],
    )
    return invoice, diagnostics


def process_file(
    path: Path,
    dpi: int,
    ocr_mode: str,
    cache_dir: Path | None,
    all_pages: bool,
) -> tuple[Invoice, SourceDiagnostics]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return process_pdf(
            path,
            dpi=dpi,
            ocr_mode=ocr_mode,
            cache_dir=cache_dir,
            all_pages=all_pages,
        )
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return process_image(path, ocr_mode=ocr_mode, cache_dir=cache_dir)
    if suffix == ".txt":
        return process_txt(path)
    if suffix == ".docx":
        return process_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def style_sheet(sheet, widths: dict[str, int]) -> None:
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def write_workbook(
    invoices: list[Invoice],
    diagnostics: list[SourceDiagnostics],
    output_path: Path,
) -> None:
    workbook = Workbook()
    items_sheet = workbook.active
    items_sheet.title = "Line Items"
    items_sheet.append(
        [
            "Source File",
            "Page",
            "Vendor Name",
            "Invoice Date",
            "Invoice Number",
            "Ship To Address",
            "Item Number",
            "Description",
            "Quantity",
            "Unit Price",
            "Unit",
            "Line Total",
            "Confidence",
            "Extraction Method",
        ]
    )

    summary_sheet = workbook.create_sheet("Invoice Summary")
    summary_sheet.append(
        [
            "Source File",
            "Vendor Name",
            "Invoice Date",
            "Invoice Number",
            "Ship To Address",
            "Line Item Count",
            "Notes",
        ]
    )

    review_sheet = workbook.create_sheet("Needs Review")
    review_sheet.append(
        [
            "Source File",
            "Page",
            "Problem",
            "Vendor Name",
            "Invoice Date",
            "Item Number",
            "Description",
            "Quantity",
            "Unit Price",
            "Unit",
        ]
    )

    raw_sheet = workbook.create_sheet("Raw Text")
    raw_sheet.append(["Source File", "Page", "Extracted Text"])

    diag_sheet = workbook.create_sheet("Diagnostics")
    diag_sheet.append(
        [
            "Source File",
            "Source Type",
            "Pages",
            "OCR Pages",
            "Avg OCR Confidence",
            "Notes",
        ]
    )

    diagnostic_map = {item.source_file: item for item in diagnostics}

    for invoice in invoices:
        summary_sheet.append(
            [
                invoice.source_file,
                invoice.vendor_name,
                invoice.invoice_date,
                invoice.invoice_number,
                invoice.ship_to_address,
                len(invoice.items),
                "; ".join(dict.fromkeys(invoice.notes)),
            ]
        )
        for page_number, text in enumerate(invoice.raw_pages, start=1):
            raw_sheet.append([invoice.source_file, page_number, text])

        if not invoice.items:
            review_sheet.append(
                [
                    invoice.source_file,
                    "",
                    "No line items found",
                    invoice.vendor_name,
                    invoice.invoice_date,
                    "",
                    "",
                    "",
                    "",
                    "",
                ]
            )

        for item in invoice.items:
            items_sheet.append(
                [
                    invoice.source_file,
                    item.page,
                    invoice.vendor_name,
                    invoice.invoice_date,
                    invoice.invoice_number,
                    invoice.ship_to_address,
                    item.item_number,
                    item.description,
                    item.quantity,
                    item.unit_price,
                    item.unit,
                    item.line_total,
                    item.confidence,
                    item.extraction_method,
                ]
            )
            problems = []
            if item.confidence == "Review":
                problems.append("Check extracted line item")
            if not item.description:
                problems.append("Missing description")
            if item.unit_price is None and item.line_total is None:
                problems.append("Missing price/amount")
            if problems:
                review_sheet.append(
                    [
                        invoice.source_file,
                        item.page,
                        "; ".join(problems),
                        invoice.vendor_name,
                        invoice.invoice_date,
                        item.item_number,
                        item.description,
                        item.quantity,
                        item.unit_price,
                        item.unit,
                    ]
                )

        for note in dict.fromkeys(invoice.notes):
            if not invoice.items and note.lower() == "no line items found":
                continue
            if any(
                keyword in note.lower()
                for keyword in (
                    "not found",
                    "no line items",
                    "low ocr",
                    "handwriting",
                    "failed",
                    "review",
                )
            ):
                review_sheet.append(
                    [
                        invoice.source_file,
                        "",
                        note,
                        invoice.vendor_name,
                        invoice.invoice_date,
                        "",
                        "",
                        "",
                        "",
                        "",
                    ]
                )

        diag = diagnostic_map.get(invoice.source_file)
        if diag:
            diag_sheet.append(
                [
                    diag.source_file,
                    diag.source_type,
                    diag.pages,
                    diag.ocr_pages,
                    diag.avg_ocr_confidence,
                    "; ".join(dict.fromkeys(diag.notes)),
                ]
            )

    currency_format = '$#,##0.00'
    for row in range(2, items_sheet.max_row + 1):
        items_sheet.cell(row, 10).number_format = currency_format
        items_sheet.cell(row, 12).number_format = currency_format
    for row in range(2, review_sheet.max_row + 1):
        review_sheet.cell(row, 9).number_format = currency_format

    style_sheet(
        items_sheet,
        {
            "A": 32,
            "B": 8,
            "C": 28,
            "D": 14,
            "E": 18,
            "F": 48,
            "G": 18,
            "H": 50,
            "I": 12,
            "J": 14,
            "K": 10,
            "L": 14,
            "M": 12,
            "N": 20,
        },
    )
    style_sheet(
        summary_sheet,
        {"A": 32, "B": 28, "C": 14, "D": 18, "E": 52, "F": 15, "G": 60},
    )
    style_sheet(
        review_sheet,
        {
            "A": 32,
            "B": 8,
            "C": 46,
            "D": 28,
            "E": 14,
            "F": 18,
            "G": 50,
            "H": 12,
            "I": 14,
            "J": 10,
        },
    )
    style_sheet(raw_sheet, {"A": 32, "B": 8, "C": 120})
    style_sheet(diag_sheet, {"A": 32, "B": 14, "C": 8, "D": 10, "E": 18, "F": 80})
    for row in range(2, raw_sheet.max_row + 1):
        raw_sheet.row_dimensions[row].height = 90
    for row in range(2, diag_sheet.max_row + 1):
        diag_sheet.row_dimensions[row].height = 45

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract invoice fields locally from PDFs, images, TXT, and DOCX files."
    )
    parser.add_argument("input", help="File or folder to process.")
    parser.add_argument(
        "-o",
        "--output",
        default="invoices_local_robust.xlsx",
        help="Output Excel workbook. Default: invoices_local_robust.xlsx",
    )
    parser.add_argument(
        "--dpi",
        type=int,
        default=200,
        help="PDF OCR render DPI. Higher is slower but may improve scans. Default: 200",
    )
    parser.add_argument(
        "--ocr-mode",
        choices=("fast", "balanced", "deep"),
        default="balanced",
        help="OCR effort. fast=one pass, balanced=retry only weak pages, deep=more retries only when needed. Default: balanced",
    )
    parser.add_argument(
        "--deep-ocr",
        action="store_true",
        help="Shortcut for --ocr-mode deep. Deep mode is selective now, not every pass on every page.",
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=0,
        help="Files to process at the same time. Default: up to 4, based on CPU count.",
    )
    parser.add_argument(
        "--all-pages",
        action="store_true",
        help="Force OCR on every scanned page. Slower. By default, attachment pages are skipped after invoice data is found.",
    )
    parser.add_argument(
        "--no-cache",
        action="store_true",
        help="Disable OCR cache. By default OCR text is cached so reruns are much faster.",
    )
    parser.add_argument(
        "--cache-dir",
        default="",
        help="Optional OCR cache folder. Default: .invoice_ocr_cache next to the output file.",
    )
    parser.add_argument(
        "--max-files",
        type=int,
        default=0,
        help="Optional limit for testing. Example: --max-files 5",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if EXTRACTOR_RULES_VERSION != REQUIRED_RULES_VERSION:
        print(
            "Error: parser version mismatch inside process_invoices_onefile.py. Re-download the latest standalone file.",
            file=sys.stderr,
        )
        return 2
    try:
        files = find_files(Path(args.input).expanduser().resolve())
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 2

    if args.max_files:
        files = files[: args.max_files]
    if not files:
        print("No supported files found.", file=sys.stderr)
        return 2

    invoices: list[Invoice] = []
    diagnostics: list[SourceDiagnostics] = []
    failures = 0
    ocr_mode = "deep" if args.deep_ocr else args.ocr_mode
    output_path = Path(args.output).expanduser().resolve()
    cache_dir = None
    if not args.no_cache:
        cache_dir = (
            Path(args.cache_dir).expanduser().resolve()
            if args.cache_dir
            else output_path.parent / ".invoice_ocr_cache"
        )
    workers = args.workers or min(4, max(1, (os.cpu_count() or 2) - 1))
    workers = max(1, workers)

    print(f"Found {len(files)} supported file(s).")

    def run_one(path: Path) -> tuple[Path, Invoice, SourceDiagnostics, Exception | None]:
        try:
            invoice, diag = process_file(
                path,
                dpi=args.dpi,
                ocr_mode=ocr_mode,
                cache_dir=cache_dir,
                all_pages=args.all_pages,
            )
            return path, invoice, diag, None
        except Exception as exc:
            invoice = Invoice(
                source_file=path.name,
                notes=[f"Processing failed: {exc}", "No line items found"],
            )
            diag = SourceDiagnostics(
                source_file=path.name,
                source_type=path.suffix.lower().lstrip(".").upper(),
                notes=[f"Processing failed: {exc}"],
            )
            return path, invoice, diag, exc

    if workers == 1 or len(files) == 1:
        for index, path in enumerate(files, start=1):
            print(f"[{index}/{len(files)}] Processing {path.name}...")
            path, invoice, diag, exc = run_one(path)
            if exc:
                failures += 1
                print(f"  WARNING: {exc}", file=sys.stderr)
            invoices.append(invoice)
            diagnostics.append(diag)
    else:
        print(f"Using {workers} worker(s), OCR mode: {ocr_mode}, DPI: {args.dpi}.")
        completed_count = 0
        results_by_index: dict[int, tuple[Invoice, SourceDiagnostics]] = {}
        with ThreadPoolExecutor(max_workers=workers) as executor:
            future_map = {
                executor.submit(run_one, path): (index, path)
                for index, path in enumerate(files)
            }
            for future in as_completed(future_map):
                index, path = future_map[future]
                completed_count += 1
                _, invoice, diag, exc = future.result()
                if exc:
                    failures += 1
                    print(f"[{completed_count}/{len(files)}] WARNING {path.name}: {exc}", file=sys.stderr)
                else:
                    print(f"[{completed_count}/{len(files)}] Done {path.name}")
                results_by_index[index] = (invoice, diag)
        for index in range(len(files)):
            invoice, diag = results_by_index[index]
            invoices.append(invoice)
            diagnostics.append(diag)

    write_workbook(invoices, diagnostics, output_path)

    review_items = sum(
        1
        for invoice in invoices
        for item in invoice.items
        if item.confidence == "Review"
    )
    print(f"\nDone: {output_path}")
    print(f"Processed: {len(invoices)} | Failed: {failures} | Review line items: {review_items}")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
